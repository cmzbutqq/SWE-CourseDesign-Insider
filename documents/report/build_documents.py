from __future__ import annotations

import shutil
import unicodedata
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
BASE_DIR = ROOT / "documents" / "base"
REPORT_DIR = ROOT / "documents" / "report"

TITLE = "一体化监控平台开发"
LEADER = "陈梦泽"
MEMBERS = "陈紫暄、曹丹"
MENTOR = "苏锦钿、黄平"
COLLEGE = "计算机科学与工程学院"
MAJOR = "计算机科学与技术"
COURSE_CODE = "045102191"
CREDITS = "2.0"
START_DATE = "2026年3月"
REPO_URL = "https://github.com/cmzbutqq/SWE-CourseDesign-Insider"
DEPLOY_PUBLIC_URL = "http://103.236.55.76/"
DEPLOY_EXPIRES = "2026年07月14日"
COVER_TITLE = " 软件工程课程设计报告书"
BODY_FONT = "宋体"
TITLE_FONT = "黑体"
TABLE_BORDER_COLOR = "9EA7B3"
TABLE_HEADER_FILL = "E6EBF2"
TABLE_ALT_FILL = "F7F9FC"
IMAGE_MAX_PAGE_HEIGHT_RATIO = 0.6


def para(text: str) -> dict:
    return {"type": "paragraph", "text": text}


def image(path: str, width: float, caption: str) -> dict:
    return {"type": "image", "path": path, "width": width, "caption": caption}


def table(rows: list[list[str]], widths: list[float] | None = None) -> dict:
    return {"type": "table", "rows": rows, "widths": widths}


def set_run_style(
    run,
    *,
    font_name: str,
    size_pt: float,
    bold: bool | None = None,
    underline: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.underline = underline
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def display_width(text: str) -> int:
    return sum(2 if unicodedata.east_asian_width(char) in {"W", "F"} else 1 for char in text)


def padded_cover_value(value: str, target_width: int) -> str:
    text = f" {value}"
    padding = max(target_width - display_width(text), 1)
    return text + (" " * padding)


def padded_text_for_runs(text: str, runs) -> str:
    slot_total = sum(len(run.text or "") for run in runs)
    if len(text) >= slot_total:
        return text
    return text + (" " * (slot_total - len(text)))


def split_text_by_lengths(text: str, lengths: list[int]) -> list[str]:
    if not lengths:
        return []
    pieces: list[str] = []
    cursor = 0
    for index, length in enumerate(lengths):
        if index == len(lengths) - 1:
            pieces.append(text[cursor:])
        else:
            pieces.append(text[cursor:cursor + length])
            cursor += length
    return pieces


def rewrite_runs_preserving_structure(runs, text: str) -> None:
    lengths = [len(run.text or "") for run in runs]
    if not lengths:
        return
    pieces = split_text_by_lengths(text, lengths)
    for run, piece in zip(runs, pieces):
        run.text = piece


def set_cover_line(paragraph: Paragraph, text: str) -> None:
    rewrite_runs_preserving_structure(paragraph.runs, text)


def set_cover_field(paragraph: Paragraph, label: str, value: str, target_width: int) -> None:
    label_runs = [run for run in paragraph.runs if not run.underline]
    underline_runs = [run for run in paragraph.runs if run.underline]
    rewrite_runs_preserving_structure(label_runs, padded_text_for_runs(label, label_runs))
    padded_value = padded_text_for_runs(f" {value}", underline_runs)
    rewrite_runs_preserving_structure(underline_runs, padded_value)


def set_cell_margins(cell, *, top: int = 90, bottom: int = 90, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if fill is None:
        if shd is not None:
            tc_pr.remove(shd)
        return
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_table_borders(table_obj: Table, color: str = TABLE_BORDER_COLOR, size: int = 6) -> None:
    tbl_pr = table_obj._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def page_height_inches(document: Document) -> float:
    section = document.sections[0]
    return section.page_height.inches


def usable_page_width_inches(document: Document) -> float:
    section = document.sections[0]
    return section.page_width.inches - section.left_margin.inches - section.right_margin.inches


def picture_kwargs(document: Document, image_path: Path, width_in: float) -> dict:
    with Image.open(image_path) as img:
        pixel_width, pixel_height = img.size
    default_height = width_in * pixel_height / pixel_width
    max_height = page_height_inches(document) * IMAGE_MAX_PAGE_HEIGHT_RATIO
    if default_height > max_height:
        return {"height": Inches(max_height)}
    return {"width": Inches(width_in)}


def normalize_weights(weights: list[float]) -> list[float]:
    positive = [max(weight, 0.0) for weight in weights]
    total = sum(positive)
    if total <= 0:
        return [1 / len(weights)] * len(weights)
    return [weight / total for weight in positive]


def table_min_column_width(col_count: int) -> float:
    if col_count <= 2:
        return 1.2
    if col_count == 3:
        return 1.0
    if col_count == 4:
        return 0.86
    return 0.72


def resolve_column_widths(document: Document, rows: list[list[str]], width_hints: list[float] | None = None) -> list[float]:
    col_count = len(rows[0])
    if width_hints and len(width_hints) >= col_count:
        weights = [float(value) for value in width_hints[:col_count]]
    else:
        weights = []
        for col_index in range(col_count):
            max_width = max(display_width(row[col_index].replace("`", "")) for row in rows)
            weights.append(min(max(max_width, 8), 28))

    total_width = usable_page_width_inches(document) - 0.2
    min_width = table_min_column_width(col_count)
    widths = [total_width * ratio for ratio in normalize_weights(weights)]
    widths = [max(width, min_width) for width in widths]
    overflow = sum(widths) - total_width
    if overflow > 0:
        reducible = [max(width - min_width, 0) for width in widths]
        reducible_total = sum(reducible)
        if reducible_total > 0:
            widths = [
                width - overflow * reducible[index] / reducible_total
                for index, width in enumerate(widths)
            ]
    return widths


def apply_column_widths(table_obj: Table, widths: list[float]) -> None:
    for col_index, width_in in enumerate(widths):
        width = Inches(width_in)
        if col_index < len(table_obj.columns):
            table_obj.columns[col_index].width = width
        for cell in table_obj.columns[col_index].cells:
            cell.width = width


def set_paragraph_text_like_existing(paragraph: Paragraph, text: str) -> None:
    rewrite_runs_preserving_structure(paragraph.runs, text)


REPORT_CONTENT = {
    "1、引言": [
        table(
            [
                ["项目项", "内容"],
                ["题目", TITLE],
                ["小组成员", f"{LEADER}（组长）；{MEMBERS}"],
                [
                    "人员分工",
                    "陈梦泽负责整体架构、Agent、后端、部署、CI/CD 和答辩主线；陈紫暄负责首页总览、趋势、节点筛选和节点详情增强；曹丹负责服务视图导航化、趋势图信息补强和视频后期。",
                ],
                [
                    "主要交付",
                    "统一监控工作台、Go Agent、Spring Boot 聚合后端、Docker Compose 演示环境、Prometheus/Grafana/SkyWalking 集成、自动化脚本、答辩 PPT 和演示视频。",
                ],
                [
                    "当前完成情况",
                    "已经形成可运行闭环，覆盖节点纳管、服务发现、主机资源采集、统一总览、服务下钻、链路追踪、部署验收和基础自动化验证。",
                ],
                [
                    "仓库与部署入口",
                    f"GitHub 仓库：{REPO_URL}；公网部署地址：{DEPLOY_PUBLIC_URL}；当前云服务器到期日期：{DEPLOY_EXPIRES}。",
                ],
            ],
            [1.35, 4.8],
        ),
        image("pics/总览.png", 6.2, "统一工作台总览"),
    ],
    "1.1 项目背景": [
        para("课题来源于中国电信广东公司 2026 年实训课题二。原始要求很直接：在主机侧部署 Agent，自动识别服务和组件，把主机信息、资源利用率、技术栈、指标和调用链放到一个监控台里展示。"),
        para("按当前实现口径，平台统一收拢了三类信息：节点主机名、IP、操作系统、Agent 版本等基础信息；CPU、内存、磁盘、网络等主机资源利用率；服务目录、指标抓取入口和业务调用链等服务侧信息。这三类信息与题目原始要求是一一对应的。"),
        para("仓库当前的实现把自研 Agent、聚合后端、前端工作台和 Prometheus、Grafana、SkyWalking 串成了一套课程级闭环系统。数据来源、排障路径和系统边界都能在当前代码、脚本和页面中对应起来。"),
        para("从课程定位看，这个题目考察的不是单一语言编码能力，而是软件工程全流程能力：需求拆解、模块划分、接口设计、容器化运行、自动化验证和团队协作都要落到一套真实项目里。当前仓库的实现路径与这一要求是吻合的。"),
    ],
    "1.2 项目目的": [
        para("本项目的目标是做出一套可演示、可联调、可部署的一体化监控平台。平台既要能看到节点和服务，也要能追到一条真实业务调用链。"),
        para("项目重点不只是页面展示，还包括把 Agent 上报、后端聚合、观测底座和前台工作台接通，让同一条业务链路既能在统一工作台中看到，也能回到原生后台继续核对。"),
        para("节点纳管和服务发现已经做成自动化过程；主机资源、服务指标和链路数据统一进入同一入口；总览、下钻和调试三类使用路径保持连通；测试、部署和演示过程可以重复执行。"),
        para("交付结果不仅包括页面，还包括 Compose 编排、GitHub Actions 工作流、部署脚本、冒烟脚本、答辩截图和演示视频。这些内容共同组成课程项目的完整提交物。"),
        table(
            [
                ["课题要求", "当前落地方式", "当前验收入口"],
                ["Agent 部署到主机", "`app-node` 与 `middleware-node` 镜像内置 `monitor-agent`。", "节点列表、部署脚本、smoke test。"],
                ["自动扫描与服务发现", "Agent 扫描 `ps` 与 `ss -ltnp`，识别固定服务类型。", "`/api/services`、服务页。"],
                ["展示主机信息与利用率", "心跳上报 CPU、内存、磁盘、网络指标，节点页与节点详情页展示。", "`/api/nodes`、`/api/nodes/{id}`。"],
                ["展示服务指标与调用链", "Prometheus 抓取指标，SkyWalking 采集 Trace，前台统一入口展示。", "服务详情、链路页、调试页。"],
                ["自动化验收", "GitHub Actions 运行基础测试、配置回归、完整冒烟和远端部署。", "`.github/workflows/ci.yml`、`deploy-main.yml`。"],
            ],
            [1.35, 2.9, 2.0],
        ),
    ],
    "1.3 项目特色": [
        para("第一，统一入口清楚。前台按总览、节点、服务、链路、调试五个视角组织页面，用户不需要在多个后台之间来回切换。"),
        para("统一入口不只是把菜单放到同一页面里。当前使用路径通常是先在总览页看节点计数、异常入口和趋势，再进入节点详情或服务详情确认资源和抓取入口，接着去链路页查看业务 Trace，最后在调试页或原生后台回证 targets 和原始查询。"),
        para("第二，演示链路是真实的。`sample-service` 会调用 `middleware-service`，后者继续访问 MySQL、Redis 和 Nginx，因此指标和 Trace 不是静态样例。"),
        para("第三，工程化比较完整。仓库内已经提供 Docker Compose 启动方式、GitHub Actions 测试与部署工作流、配置回归脚本和冒烟脚本，适合课程展示和复现。"),
        para("仓库当前还启用了 `main` 分支保护：禁止删除和强推，改动需通过 Pull Request 合入，至少 1 个评审通过并解决评审线程，`frontend-build`、`backend-test`、`agent-test` 三项状态检查必须通过，并要求按最新主分支严格复核。"),
        para("第四，平台后端除保存原始节点和服务信息外，还生成总览摘要、异常入口、趋势快照、快捷跳转链接和 Trace 摘要。前台直接读取的是聚合后的排障信息。"),
    ],
    "1.4 项目风险": [
        table(
            [
                ["风险点", "当前表现", "当前处理"],
                ["异常识别演示场景", "MySQL 被故意保留为无独立 exporter 的状态，因此缺少 metricsPath。", "前台直接标成 `NO_SCRAPE_ENDPOINT`，用于证明平台能够识别缺少抓取入口的异常服务。"],
                ["发现规则固定", "服务识别依赖进程名和端口规则，当前覆盖面是固定集合。", "规则已经覆盖 Spring Boot、MySQL、Redis、Nginx、node_exporter 等课程展示需要的对象。"],
                ["运行环境偏演示", "前端与后端容器使用开发态启动方式，SkyWalking OAP 用 H2。", "当前目标是课程演示与联调闭环，不把这一版描述成生产部署。"],
                ["业务 Trace 依赖流量", "没有流量时，SkyWalking 可能看不到业务 Trace。", "app-node 内置 `traffic-generator`，部署脚本也会主动 warmup `demo-chain`。"],
                ["统一入口依赖代理路径", "Grafana、Prometheus 和前台代理路径一旦改错，页面容易出现嵌入失败或跳转失效。", "仓库中保留 route-prefix 与 observability 配置回归脚本，优先通过自动化脚本拦截这类回归。"],
            ],
            [1.2, 2.05, 2.95],
        ),
        para("这些风险主要影响课程演示边界说明，不构成平台主链路失效。"),
    ],
    "1.5 产品范围": [
        para("本项目当前范围包括：节点注册、节点心跳、主机资源采集、规则式服务发现、统一前台工作台、Grafana/Prometheus/SkyWalking 嵌入与跳转、演示链路、部署脚本、配置回归脚本和冒烟脚本。"),
        para("不在当前范围内的部分也要说清楚：它不是一套通用服务发现平台，不负责自动改写 Prometheus 抓取配置，也没有把前后端打磨成生产级长期运维形态。"),
        para("对应到具体页面，产品范围已经覆盖总览、节点、节点详情、服务、服务详情、链路追踪和调试面板七类核心展示入口。对应到后端能力，已经覆盖注册、心跳、聚合读接口、趋势快照和 Trace 摘要五类核心服务。"),
        table(
            [
                ["监控对象", "识别方式", "当前采集或接入方式", "当前指标入口或状态"],
                ["Linux 主机", "Agent 所在节点自动纳管", "Agent 读取 /proc、df 等信息并随心跳上报", "node_exporter /metrics"],
                ["Spring Boot 应用", "java 进程与 Spring 启动特征", "Actuator + Micrometer", "/actuator/prometheus"],
                ["Nginx", "进程名含 nginx", "nginx exporter 抓取 nginx_status", ":9113/metrics"],
                ["Redis", "进程名含 redis-server", "redis_exporter", ":9121/metrics"],
                ["MySQL", "进程名含 mysqld 或 mariadbd", "仅做服务识别", "刻意不接 exporter，前台标为 NO_SCRAPE_ENDPOINT"],
                ["node_exporter", "进程名含 node_exporter", "Prometheus 直接抓取", "/metrics"],
            ],
            [1.05, 1.35, 2.0, 2.45],
        ),
    ],
    "2、综合描述": [
        para("系统从结构上可以分成五层：Agent 负责采集与上报；后端负责聚合与判断；MySQL 负责平台元数据；Prometheus、Grafana、SkyWalking 提供指标和链路能力；前端作为统一入口展示与回跳。"),
        table(
            [
                ["核心组件", "部署位置或 profile", "主要端口或入口", "主要作用"],
                ["frontend", "核心开发层", "15173 / `/overview`", "统一工作台入口。"],
                ["backend", "核心开发层", "18081 / `/api/overview`", "聚合节点、服务、趋势和 Trace 摘要。"],
                ["mysql", "核心开发层", "13306", "保存平台元数据、资源指标和平台快照。"],
                ["prometheus", "observability", "19090 / `/prometheus`", "抓取应用指标和 exporter 指标。"],
                ["grafana", "observability", "13000 / `/grafana`", "展示平台总览、节点和服务仪表盘。"],
                ["skywalking-oap / skywalking-ui", "observability", "11800 / 12800 / 18082", "接收与展示业务 Trace。"],
                ["app-node", "nodes", "8081 / 9100", "承载 `sample-service`、node_exporter、Agent 和持续造流量组件。"],
                ["middleware-node", "nodes", "8082 / 3306 / 6379 / 9113 / 9121", "承载 `middleware-service`、MySQL、Redis、Nginx 和相关 exporter。"],
            ],
            [1.4, 1.35, 1.6, 2.1],
        ),
        image("pics/腾讯文档-总体架构图.png", 5.6, "平台总体架构图"),
        para("平台控制面、观测底座和语义节点在这张图里被放到同一张拓扑上。Agent 向后端上报，Prometheus 面向指标入口抓取，SkyWalking 面向 Java Agent 接收链路数据，前端负责把这些结果组织成统一工作台。"),
        image("pics/服务视图.png", 5.8, "服务视图总览"),
    ],
    "2.1 用户类和特性": [
        para("平台使用者包括课程答辩教师、平台演示人员和开发维护成员。不同角色关注点不同，常用入口也不同。"),
        table(
            [
                ["用户类", "关注点", "典型使用方式"],
                ["课程答辩教师", "是否满足课题要求，系统是否真能跑通。", "先看总览页，再看节点、服务、链路和自动化材料。"],
                ["平台演示人员", "能否快速说明问题、展示数据来源、顺着异常往下钻。", "使用总览、节点详情、服务详情和链路页完成演示。"],
                ["开发维护成员", "是否方便定位配置问题、抓取问题和链路问题。", "从调试页、Prometheus targets、Grafana 面板和 SkyWalking Trace 继续排查。"],
            ],
            [1.25, 2.2, 2.65],
        ),
        para("三类角色的使用顺序也比较清楚：教师先看总览、节点、服务和链路页确认课题完成度；演示人员主要依赖总览、节点详情和服务详情组织讲解；维护成员更多使用调试页、Prometheus targets 和 SkyWalking 原生界面。"),
    ],
    "2.2 运行环境": [
        para("项目按核心开发层、观测层和演示节点层拆成不同 profile。日常联调和答辩演示使用的启动范围不同。"),
        para(f"当前代码仓库地址为 {REPO_URL}。当前公网部署访问地址为 {DEPLOY_PUBLIC_URL}，该云服务器租期到 {DEPLOY_EXPIRES}，可作为答辩与验收期间的在线演示入口。"),
        para("当前仓库中的关键运行版本可以直接核对到代码与编排文件：后端使用 Spring Boot 3.3.5 和 Java 21，前端依赖 Vue 3.5.13、Vue Router 4.5.0、ECharts 5.6.0、Vite 8.0.14，容器使用 Node 22，Agent 使用 Go 1.23，观测组件使用 Prometheus 2.54.1、Grafana 11.2.2、SkyWalking 10.0.1，数据库为 MySQL 8.4。"),
        para("Agent 采集依赖 `/proc/stat`、`/proc/meminfo`、`/proc/net/dev` 以及 `df`、`ps`、`ss` 等 Linux 侧能力，因此被纳管节点当前默认假定为 Linux 环境。"),
        table(
            [
                ["类别", "内容"],
                ["运行底座", "Docker Compose；核心层为 mysql、backend、frontend；观测层为 prometheus、grafana、skywalking-oap、skywalking-ui；演示节点层为 app-node、middleware-node。"],
                ["前端", "Vue 3 + Vite，默认入口 `http://localhost:15173/overview`。"],
                ["后端", "Spring Boot 3.3.5、Java 21、JPA、MySQL，默认 API 入口 `http://localhost:18081/api/overview`。"],
                ["Agent", "Go 1.23；注册间隔 15 秒，心跳间隔 10 秒。"],
                ["观测组件", "Prometheus 15 秒抓取；Grafana 负责面板；SkyWalking 负责 Trace。"],
                ["演示条件", "课题原始建议为 2 台 4C8G 主机；当前实现用 1 台 8C16G 云服务器或本地 Docker 环境模拟 2 个语义节点。"],
            ],
            [1.35, 4.8],
        ),
        image("pics/节点视图.png", 5.8, "节点视图总览"),
        para(f"本报告中的前台页面截图均来自公网部署地址 {DEPLOY_PUBLIC_URL} 的实际访问结果，不使用本地开发环境截图。"),
    ],
    "2.3 设计和实现上的限制": [
        para("服务识别由 `agent/internal/discovery/discovery.go` 中的固定规则完成，当前只内置 `sample-service`、`middleware-service`、`nginx`、`mysql`、`redis` 和 `node-exporter` 六类对象的识别规则。"),
        para("Prometheus 抓取目标通过 `infra/prometheus/prometheus.yml` 静态配置。当前抓取对象包括 backend、两台语义节点、两套 Spring Boot 服务、nginx 和 redis，不包含 MySQL。Agent 发现到的新服务会先进入平台服务目录，不会自动写回抓取配置。"),
        para("前端容器执行 `npm run dev -- --host 0.0.0.0 --port 5173`，后端通过 `spring-boot:run` 启动，SkyWalking OAP 使用 `SW_STORAGE=h2`。当前部署形态用于课程演示和联调。"),
        para("Agent 能识别 MySQL，但当前未配置独立 exporter，因此服务页会把它标记为 `NO_SCRAPE_ENDPOINT`。"),
        para("节点列表当前通过 `managedNodeRepository.findAll().stream()` 在内存中按状态、关键字和服务类型过滤，再做排序。这说明平台规模假设仍偏向课程设计和小规模节点场景。"),
        table(
            [
                ["限制项", "仓库依据", "当前影响"],
                ["服务发现规则固定", "`discovery.go` 只覆盖五类课程演示对象。", "未命中的服务不会自动进入统一工作台。"],
                ["Prometheus targets 静态配置", "`docker-compose.yml` 挂载 `infra/prometheus/prometheus.yml`。", "自动发现不等于自动抓取，新对象仍需补配置。"],
                ["控制面以开发态运行", "前端走 Vite dev server，后端走 `spring-boot:run`。", "当前未实现多实例、高可用和审计。"],
                ["SkyWalking 存储为 H2", "`docker-compose.yml` 中 `SW_STORAGE=h2`。", "适合课程场景，不用于长期高并发链路留存。"],
                ["MySQL 故意不接 exporter", "Agent 可识别 MySQL，但服务对象没有 `metricsPath`。", "前台会把 MySQL 归为异常服务。"],
            ],
            [1.35, 2.45, 2.2],
        ),
    ],
    "2.4 假设和约束（依赖）": [
        para("平台运行依赖 Docker、Docker Compose 和容器网络。核心层、观测层和演示节点层都建立在同一套 Compose 编排上。"),
        para("链路页依赖真实业务流量。仓库通过 `traffic-generator.sh` 和部署脚本中的 `warmup_business_traces` 持续触发 `/api/demo-chain`，用来稳定产出业务 Trace。"),
        para("远端自动部署依赖目标服务器预置 `.env`、Docker 环境、Git 仓库和 SSH 凭据；GitHub Actions 还依赖 `DEPLOY_HOST`、`DEPLOY_PORT`、`DEPLOY_USER`、`DEPLOY_SSH_KEY` 等 secrets。"),
        para("当前两台语义节点都以 Compose 容器方式接入同一网络，便于 Prometheus 抓取、后端互访和演示链路贯通。如果只启动核心层而不补 observability 与 nodes profile，工作台虽然能打开，但不会形成完整的指标和业务链路展示。"),
        table(
            [
                ["依赖项", "主要作用", "缺失后的直接结果", "当前验证入口"],
                ["Docker 与 Compose", "拉起前端、后端、观测组件和演示节点。", "整套平台无法启动。", "`README.md`、smoke test。"],
                ["容器网络互通", "保证 Prometheus、Grafana、SkyWalking、Agent 和后端互访。", "页面可开但数据链路不完整。", "`docker-compose.yml`、部署脚本健康检查。"],
                ["持续业务流量", "让 SkyWalking 产生业务 Trace。", "链路页只能看到组件可用，看不到业务链。", "`traffic-generator.sh`、`wait_for_business_traces`。"],
                ["服务器 `.env` 与 SSH 凭据", "支持远端拉取代码、保留配置并执行部署。", "工作流只能完成仓库侧测试，不能完成服务器部署。", "`deploy-main.yml`、`deploy-main.sh`。"],
                ["观测代理与 route-prefix 配置", "保证 `/prometheus`、`/grafana`、SkyWalking 入口可嵌入可跳转。", "调试页和快捷链接容易失效。", "`prometheus-route-prefix-test.sh`。"],
            ],
            [1.25, 1.7, 2.0, 1.85],
        ),
    ],
    "3、功能性需求分析": [
        para("功能需求覆盖节点接入、状态采集、统一展示、服务下钻、链路追踪和部署验收。"),
        image("pics/腾讯文档-系统用例图.png", 5.6, "系统主要用例图"),
        para("系统主要用例集中在四类动作：接入节点、查看平台态势、下钻节点与服务、确认业务链路。四类动作与页面、接口和后台截图一一对应。"),
        table(
            [
                ["功能模块", "输入或依赖", "输出结果"],
                ["节点注册与服务发现", "Agent 读取环境变量，扫描 `ps` 与 `ss -ltnp`。", "后端创建或更新节点，保存 Spring Boot、MySQL、Redis、Nginx、node_exporter 等服务目录。"],
                ["节点心跳与资源采集", "Agent 采集 CPU、内存、磁盘、网络数据并定时上报。", "后端刷新状态、保存心跳事件和节点指标，前台节点页与节点详情页可见。"],
                ["统一总览", "后端汇总节点、服务、异常、趋势与外链信息。", "前台总览页展示风险摘要、异常入口、趋势图和多源面板。"],
                ["节点与服务下钻", "前台调用 `/api/nodes`、`/api/services` 及详情接口。", "用户可从节点、服务两个方向查看资源指标、基础信息和原始监控入口。"],
                ["链路追踪与调试", "SkyWalking 接收 Java Agent Trace；后端通过 GraphQL 查询摘要。", "前台链路页先展示最近业务 Trace 摘要，再引导进入 SkyWalking 原视图。"],
                ["部署与验收", "GitHub Actions、部署脚本、配置回归脚本、冒烟脚本。", "平台可以在 push 到 main 后完成测试、远端部署和关键健康检查。"],
            ],
            [1.35, 2.15, 2.7],
        ),
        table(
            [
                ["页面或入口", "主要后端接口", "主要外部观测入口", "页面中的核心信息"],
                ["总览", "`/api/overview`、`/api/trends`", "Grafana、Prometheus、SkyWalking", "节点数、服务数、异常入口、趋势图、统一跳转入口。"],
                ["节点列表 / 节点详情", "`/api/nodes`、`/api/nodes/{id}`", "Grafana node-detail、Prometheus、SkyWalking", "节点状态、资源指标、服务清单、快捷跳转。"],
                ["服务列表 / 服务详情", "`/api/services`、`/api/services/{id}`", "Grafana service-detail、Prometheus、SkyWalking", "服务基础信息、指标入口、调用链入口和异常标识。"],
                ["链路页", "`/api/tracing/summary`", "SkyWalking UI", "最近业务 Trace、服务链和依赖链摘要。"],
                ["调试页", "Prometheus API 摘要", "Prometheus targets 与 graph", "抓取状态、原始 query 入口和标签信息。"],
            ],
            [1.5, 1.65, 1.75, 2.0],
        ),
        image("pics/腾讯文档-心跳上报流程图.png", 5.2, "Agent 心跳上报流程图"),
        para("心跳上报链路决定了平台状态是否可信。节点状态、资源指标和最近心跳时间都沿着这条链路进入后端，再被总览页、节点页和趋势接口消费。"),
        para("注册接口按 `nodeName` 识别节点。重复注册不会生成重复节点，而是更新该节点的基础信息，并把本次发现到的服务目录同步到后端。这保证节点名既是展示标识，也是平台当前去重和持续更新的主键。"),
        para("心跳接口除刷新状态外，还会写入 `HeartbeatEvent`；如果本次指标不完整，后端会与上一条 `NodeMetrics` 合并后再入库，尽量保持节点详情页的资源图表连续。"),
        image("pics/节点详情（资源监控）.png", 6.2, "节点详情中的资源监控"),
        para("节点详情页把 CPU、内存、磁盘、网络和服务清单放在同一入口，便于从单个节点继续下钻。"),
        image("pics/服务详情（基础信息）.png", 5.8, "服务详情中的基础信息与监控入口"),
        para("服务详情页同时展示服务基础信息、抓取入口、Trace 入口和原始监控链接，便于确认某个服务的接入状态。"),
        para("节点列表支持按状态、关键字、服务类型筛选，并支持按最近心跳排序。趋势接口单次最大查询窗口为 720 小时，前台当前提供 15 分钟、30 分钟、1 小时、2 小时、4 小时和 24 小时六档快捷范围。"),
        para("接口设计也有明显的读写分工。`/api/agents/register` 和 `/api/agents/heartbeat` 面向 Agent 写入，`/api/overview`、`/api/nodes`、`/api/services`、`/api/tracing/summary` 等接口面向工作台读取聚合结果。"),
        image("pics/架构图（观测数据流）.png", 5.8, "观测数据流"),
        image("pics/链路追踪全览.png", 6.2, "链路追踪全览"),
        image("pics/调试面板.png", 5.8, "调试面板"),
        para("调试面板保留了 Prometheus targets 和原始 query 入口，用于确认抓取链路和标签是否正确。"),
    ],
    "4、非功能需求分析": [
        para("非功能要求覆盖可演示、可复现、可定位、可维护和可扩展五个方面。"),
        table(
            [
                ["工程支撑项", "当前实现", "直接结果"],
                ["基础测试", "`ci.yml` 分别运行前端、后端、Agent 的测试与构建。", "合并前即可发现基础逻辑回归。"],
                ["配置回归", "`packaging-test` 运行 SkyWalking 来源、Prometheus route-prefix 和部署观测配置检查。", "优先拦截最容易破坏演示闭环的配置错误。"],
                ["完整环境冒烟", "`tests/smoke-test.sh` 拉起 `observability` 与 `nodes` profile。", "验证节点注册、服务发现、targets 和业务 Trace。"],
                ["远端部署", "`deploy-main.sh` 等待健康检查、节点注册、Prometheus targets 和业务 Trace。", "服务器上的部署结果不止是容器启动，而是完整闭环可用。"],
            ],
            [1.5, 2.6, 2.15],
        ),
    ],
    "4.1 性能需求": [
        para("Agent 默认每 15 秒执行一次注册、每 10 秒发送一次心跳；Prometheus 默认每 15 秒抓取一次指标；后端每 5 分钟写入一次平台快照。"),
        para("趋势接口最多查询 30 天窗口，返回的是后端汇总后的节点数、服务数和告警数序列，不等价于原始高频时序。"),
        para("总览页、节点详情页和服务详情页读取的是后端聚合接口结果，不直接在浏览器侧拼接 Prometheus 或 SkyWalking 原始数据。"),
        table(
            [
                ["采集或查询项", "当前频率或方式", "直接结果"],
                ["Agent 注册", "默认每 15 秒执行一次。", "节点基础信息和服务目录可持续刷新。"],
                ["Agent 心跳", "默认每 10 秒执行一次。", "节点状态与主机指标更新较快，答辩演示等待时间较短。"],
                ["Prometheus 抓取", "默认 15 秒一次。", "服务指标和 exporter 指标可以稳定进入 Grafana 与调试页。"],
                ["平台趋势快照", "后端每 5 分钟保存一次。", "趋势图能保留历史节奏，但不会膨胀成高频原始时序。"],
                ["趋势窗口限制", "单次最多查询 30 天。", "接口口径稳定，避免无上限历史查询。"],
            ],
            [1.45, 1.8, 3.55],
        ),
        image("pics/Grafana后台(platform overview).png", 5.7, "Grafana 平台总览面板"),
    ],
    "4.2 可靠性需求": [
        para("后端对心跳状态做合法性校验，对缺失指标做与上一条数据的合并，避免图表因为单次上报不完整而断裂。"),
        image("pics/腾讯文档-Agent心跳状态流转图.png", 5.4, "Agent 心跳状态流转图"),
        para("节点状态不是简单的在线或离线二分。后端把 `ONLINE`、`OFFLINE`、`WARNING` 作为允许状态，并以 45 秒作为延迟风险提示线、90 秒作为心跳超时阈值。超过 45 秒会提示“心跳延迟，存在超时风险”，超过 90 秒则按“心跳超时，可能失联”处理。"),
        para("应用启动后会补写历史节点缺失的 `lastHeartbeatAt`，避免升级后旧节点全部显示为“等待心跳上报”。"),
        para("两台演示节点中的核心进程都由 `supervisord` 托管，并设置 `autorestart=true`。`sample-service`、`middleware-service`、`node_exporter`、`nginx-exporter`、`redis-exporter`、`monitor-agent` 和 `traffic-generator` 出现容器内单进程退出时，会自动拉起。"),
        para("数据库容器配置了 `mysqladmin ping` 健康检查，`backend` 依赖 `service_healthy` 后再启动。部署脚本在容器启动后还会继续等待前后端、Prometheus、Grafana、SkyWalking、节点注册、Prometheus targets 和业务 Trace，就算容器已经起来，也不会把“系统未通”误判成部署成功。"),
        para("部署脚本在拉起容器后，依次检查后端、前端、Prometheus、Grafana、SkyWalking、节点注册、服务目录、Prometheus targets 和 `/api/demo-chain` 业务 Trace。"),
        para("后端每 5 分钟保存一次平台快照，每周一凌晨清理 30 天前的快照和 7 天前的节点指标。"),
        para("MySQL、Prometheus 和 Grafana 都使用 named volume 持久化。普通重启不会清空平台主数据、抓取库和仪表盘状态；只有显式执行 `docker compose down -v` 之类的清卷操作时，相关数据才会被一并移除。"),
        image("pics/Prometheus后台（targets）.png", 5.7, "Prometheus targets 状态"),
    ],
    "4.3 安全性需求": [
        para("安全边界主要在部署链路和观测后台访问。部署文档明确区分“本机到服务器”的登录密钥和“服务器到 GitHub”的拉代码密钥。"),
        para("部署脚本先备份服务器 `.env`，同步代码后再恢复配置，避免 `git reset --hard origin/main` 覆盖服务器上的运行参数。"),
        para("Grafana 允许匿名只读访问，供前台嵌入监控面板。Prometheus 调试入口保留在受控演示网络内使用。"),
        para("当前安全策略以演示可用为主。后端 `/api/**` 放开 `GET`、`POST`、`OPTIONS`，并允许跨域 `*`；前台和脚本联调更直接。"),
        para("数据库密码和部署参数通过环境变量注入；后端 Actuator 当前只暴露 `health`、`info`、`prometheus` 三个端点。结合 Grafana 的匿名 Viewer 设置。"),
    ],
    "4.4 可维护性需求": [
        para("仓库目录按前端、后端、Agent、演示节点、基础设施、部署脚本和测试脚本分开，职责边界清楚。"),
        para("CI 工作流把前端、后端、Agent 三块的测试与构建拆开，配置回归脚本单独存在，失败后可以直接定位到对应层。"),
        para("`pages.yml` 只发布 `frontend/dist` 静态产物，不等于后端、Prometheus、Grafana 和 SkyWalking 一起上线。"),
        para("后端把总览、节点、服务、趋势和 Trace 摘要收敛成独立 DTO；前端把观测入口封装在独立 service 中，页面代码不需要散落拼接第三方链接。"),
        para("定位问题的顺序通常也比较固定：页面异常先看 `frontend/src/pages` 和 `src/services`，接口口径问题先看 `NodeRegistryService` 与 `SkyWalkingQueryService`，节点纳管或服务识别问题再看 `agent/main.go` 和 `discovery.go`。"),
    ],
    "4.5 可扩展性需求": [
        para("Agent 的服务识别规则和前台面板分组都还能继续扩展。后续如果要支持更多服务类型，首先改 `discovery.go` 的分类规则，再补 Prometheus targets、Grafana 面板和服务详情页入口。"),
        para("后端已将节点、服务、趋势和 Trace 摘要拆成独立 DTO 和服务接口，新增读接口时不需要改动现有页面协议。"),
        para("当前最明确的扩展方向包括补 MySQL exporter、增加异常注入场景、补浏览器级 E2E 用例，以及让动态发现结果参与更多自动化配置流程。"),
        para("如果后续继续扩展，优先级应放在三类事情上：补更多 exporter 和服务类型，补更系统的异常注入场景，补浏览器级 E2E 或更完整的自动化配置联动。"),
    ],
    "4.6 数据一致性": [
        para("节点注册和节点心跳分两条写路径处理。注册负责节点与服务目录，心跳负责状态、心跳事件和主机指标。"),
        para("数据库侧使用 JPA 实体和事务方法维护节点、服务、心跳事件、节点指标和平台快照。趋势图的数据源来自快照表，和页面即时查询口径保持一致。"),
        para("节点列表、节点详情、服务详情和趋势页都基于同一套节点表、服务表和快照表，只是展示粒度不同。SkyWalking Trace 摘要由后端实时查询 GraphQL，不回写本地数据库。"),
        para("`metrics_snapshots` 表还对 `timestamp` 建了正序和倒序索引，便于按时间窗口读取趋势和执行历史清理。平台级趋势查询因此不需要在快照表上做全表顺序扫描。"),
        table(
            [
                ["写入或读取路径", "主要实体或来源", "消费位置", "一致性体现"],
                ["节点注册", "`ManagedNode`、`DiscoveredService`。", "总览、节点页、服务页。", "节点与服务目录来自同一次注册同步。"],
                ["节点心跳", "`HeartbeatEvent`、`NodeMetrics`、`lastHeartbeatAt`。", "节点状态、节点详情、异常摘要。", "状态、资源指标和最近心跳时间同步刷新。"],
                ["趋势快照", "`MetricsSnapshot`。", "总览趋势页。", "趋势口径直接复用在线节点、异常服务和告警统计。"],
                ["Trace 摘要", "SkyWalking GraphQL。", "链路页、服务详情快捷入口。", "链路数据与本地元数据解耦，不污染平台主数据表。"],
            ],
            [1.45, 1.7, 1.8, 2.85],
        ),
        image("pics/腾讯文档-数据库关系图.png", 5.4, "平台核心数据关系"),
    ],
    "5、项目进度安排": [
        para("项目推进与仓库提交节奏基本一致，可分为骨架搭建、总览与节点页完善、状态修正与测试补强、远端部署、观测闭环和文档整理六段。"),
        table(
            [
                ["阶段", "时间", "主要工作", "对应成果"],
                ["选题与骨架搭建", "2026-03-28 至 2026-04-04", "完成仓库初始化、Docker Compose 骨架、基础文档和前端 CI 初版。", "已完成 Compose 骨架、前端 CI 和基础文档。"],
                ["总览与节点视图完善", "2026-04-14 至 2026-05-14", "补充首页状态总览、趋势图、节点筛选、心跳排序和节点详情摘要。", "已增加首页状态总览、趋势图、节点筛选和节点详情摘要。"],
                ["状态修正与测试补强", "2026-05-24 至 2026-05-31", "修复心跳状态、空服务注册和趋势查询问题，并补后端、Agent、前端相关测试。", "已修复心跳状态、空服务注册和趋势查询问题，并补后端、Agent、前端测试。"],
                ["部署与服务视图增强", "2026-06-13 至 2026-06-14", "加入 SSH 自动部署、服务导航化展示、趋势图细节增强和云端联调。", "已接入 SSH 部署工作流、服务导航化页面和趋势图增强。"],
                ["观测闭环收口", "2026-06-20 至 2026-06-22", "接入服务详情页、观测面板嵌入、真实 demo-chain、SkyWalking 配置回归和冒烟脚本。", "已补服务详情页、观测嵌入、真实 demo-chain、SkyWalking 回归和 smoke test。"],
                ["文档与答辩材料整理", "2026-06-23 至 2026-06-27", "更新架构图、截图、PPT、视频和说明文档。", "已完成架构图、截图、PPT、视频和说明文档更新。"],
            ],
            [1.0, 1.45, 2.25, 1.5],
        ),
        image("pics/GitHub actions.png", 5.8, "GitHub Actions 运行结果"),
        image("pics/GitHub action流程详情.png", 5.9, "主分支工作流流程详情"),
        para("main 分支部署流程的顺序是 packaging-test、smoke-test、SSH 部署，以及远端健康检查和业务 Trace 验证。"),
        para("流程详情图把依赖关系展示得更直观：前端、后端、Agent 的基础测试以及 packaging-test 先完成，随后才进入完整 smoke-test，最后才执行 deploy。部署前先过完整测试和冒烟测试这一约束，在工作流图里可以直接看到。"),
        para(f"当前对外可复核的交付入口有两个：一是 GitHub 仓库 {REPO_URL}，用于查看源码、工作流和文档；二是公网部署地址 {DEPLOY_PUBLIC_URL}，用于直接访问在线系统。该云服务器到期日期为 {DEPLOY_EXPIRES}。"),
        para(f"报告、测试文档和答辩 PPT 中使用的前台页面截图，均取自公网部署地址 {DEPLOY_PUBLIC_URL} 的实际访问画面，不使用本地开发环境页面截图。Prometheus、Grafana 和 SkyWalking 的回证截图也按同一套在线环境保留。"),
        para("演示视频只作展示路径补充，顺序保持为总览先看异常，再做节点与服务下钻，随后展示 Trace，最后回到原生后台复核。"),
        para("当前进度安排与仓库交付节奏基本一致：前半段先把页面、节点状态和基础测试做稳，后半段再把云端部署、观测闭环、答辩截图和视频整理补齐。"),
        para("因此，报告中的阶段成果不是单纯按时间罗列，而是对应到当前仓库里已经可以复现的页面、接口、工作流和脚本。"),
    ],
}


REQ_CONTENT = {
    "1、引言": [
        para("需求范围限定在当前仓库已经实现并可由代码、脚本、截图回证的能力，包括节点纳管、服务发现、资源采集、统一展示、链路追踪和自动部署。"),
        para("尚未落地的设想不写成现成能力。需求范围只保留可以被源码、工作流或界面材料直接证明的事实。"),
    ],
    "1.1 编写目的": [
        para("需求内容覆盖设计、编码、测试和答辩所需的同一套系统能力。主机发现、资源采集、指标展示和调用链展示都对应可验证实现。"),
        para("页面结果与系统内部实现保持一致。页面截图、接口返回、部署脚本和自动化记录描述的是同一套事实。"),
    ],
    "1.2 项目背景": [
        para("课题原始背景来自企业运维场景：系统数量增加、架构变复杂、需要一套能够发现主机上服务并集中展示的监控工具。"),
        para("因此，本项目的需求不是单做页面，而是围绕主机纳管、资源采集、服务识别、指标展示和调用链展示建立一条能运行的主线。"),
        para("项目采用“自研控制面 + 复用成熟观测底座”的策略。Agent、平台聚合后端和统一工作台承担课程要求中的自研部分，Prometheus、Grafana 和 SkyWalking 提供底层指标与链路能力。"),
    ],
    "1.3 参考资料": [
        table(
            [
                ["资料", "用途"],
                [REPO_URL, "给出当前源码、工作流、说明文档和提交历史。"],
                ["`documents/base/广东公司提供实训课题2026年版-课题二.md`", "给出课题背景、展示内容和硬件条件。"],
                ["`docs/项目说明.md`", "说明当前实现与课题要求的映射关系。"],
                ["`README.md`", "给出当前仓库的运行方式、入口地址和测试命令。"],
                ["`docker-compose.yml` 与 `infra/` 配置", "给出系统拓扑、端口和观测组件接入方式。"],
                ["`tests/` 与 `.github/workflows/`", "给出验收脚本和自动化验证入口。"],
                ["`deploy/deploy-main.sh`", "给出远端部署后的真实健康检查顺序。"],
                ["`pics/` 与 `documents/答辩PPT.pptx`", "给出工作台、观测后台和答辩展示可回证的图像材料。"],
            ],
            [2.7, 3.45],
        ),
    ],
    "1.4 术语说明": [
        table(
            [
                ["术语", "说明"],
                ["Agent", "运行在被监控节点上的 Go 程序，负责注册、心跳、服务发现和主机资源采集。"],
                ["统一工作台", "前端工作台，包含总览、节点、服务、链路和调试五类页面。"],
                ["平台聚合后端", "Spring Boot 后端，对节点、服务、趋势和 Trace 摘要做统一加工。"],
                ["语义节点", "为答辩演示而定义的逻辑主机。当前仓库里对应 `app-node` 和 `middleware-node`。"],
                ["异常服务", "当前主要指已识别但缺少抓取入口的服务，前台用 `NO_SCRAPE_ENDPOINT` 表示。演示环境中的 MySQL 是刻意保留的示例。"],
            ],
            [1.3, 4.85],
        ),
    ],
    "1.5 产品范围": [
        para("需求范围覆盖节点纳管、服务发现、主机资源采集、统一总览、节点和服务详情、链路追踪、Prometheus/Grafana/SkyWalking 联通、自动部署和冒烟验收。"),
        para("当前不要求做生产级高可用部署，不要求做动态改写 Prometheus 配置，也不要求做浏览器级端到端自动化。"),
        table(
            [
                ["范围类别", "当前要求"],
                ["必须完成", "Agent 注册与心跳、服务发现、主机指标采集、统一工作台、链路摘要、自动化测试和自动部署。"],
                ["建议完成", "截图材料、答辩演示链路、配置回归脚本、趋势分析展示。"],
                ["当前不纳入", "多租户、权限细分、动态 Prometheus 改写、生产级高可用、复杂告警编排。"],
            ],
            [1.4, 5.1],
        ),
        table(
            [
                ["监控对象", "识别方式", "当前接入方式", "当前状态"],
                ["Linux 主机", "Agent 所在节点自动纳管", "Agent 采集系统信息与资源指标", "已实现"],
                ["Spring Boot 应用", "`java` 进程与 Spring 启动特征", "Actuator + Micrometer", "已实现"],
                ["Nginx", "进程名含 `nginx`", "`nginx-prometheus-exporter`", "已实现"],
                ["Redis", "进程名含 `redis-server`", "`redis_exporter`", "已实现"],
                ["MySQL", "进程名含 `mysqld` 或 `mariadbd`", "当前仅做服务识别", "已识别，刻意未接独立 exporter"],
                ["node_exporter", "进程名含 `node_exporter`", "Prometheus 直接抓取", "已实现"],
            ],
            [1.15, 1.55, 1.95, 2.3],
        ),
    ],
    "2、综合描述": [
        para("系统的使用路径是：Agent 上报节点和服务，后端聚合平台信息，前端按不同视角展示，Prometheus、Grafana、SkyWalking 提供底层指标和链路能力。"),
        para("典型使用顺序是：先在总览页看平台状态，再从节点或服务入口下钻，必要时进入链路页或调试页继续确认底层数据是否真实存在。"),
        image("pics/腾讯文档-总体架构图.png", 5.6, "平台总体架构图"),
        para("这套系统不是单一网页，而是采集、聚合、展示三段能力同时成立的结果。Agent、后端、前端和观测底座同时接通，才满足课题里的集中监控要求。"),
        image("pics/架构图（角色与语义拓扑）.png", 5.8, "角色与语义拓扑"),
    ],
    "2.1 用户类和特性": [
        para("三类使用者分别是课程答辩教师、演示人员和维护成员。"),
        table(
            [
                ["用户类", "特性与需求"],
                ["课程答辩教师", "更关注课题完成度、链路是否真实、数据来源是否能回证。"],
                ["演示人员", "需要一个统一入口，能在 10 分钟内讲清楚系统结构和排障流程。"],
                ["项目维护者", "需要看到节点、服务、指标、Trace 和配置回归脚本之间的关系，方便继续修改。"],
            ],
            [1.4, 4.75],
        ),
        image("pics/腾讯文档-系统用例图.png", 5.6, "系统主要用例图"),
        para("用例图把不同角色的核心动作收束成统一入口、节点巡检、服务下钻和链路核对四组路径，和当前工作台的页面分工是一致的。"),
    ],
    "2.2 运行环境": [
        para(f"当前需求落地对应的外部交付入口有两个：代码仓库为 {REPO_URL}，公网部署地址为 {DEPLOY_PUBLIC_URL}。当前云服务器租期到 {DEPLOY_EXPIRES}。"),
        table(
            [
                ["项", "内容"],
                ["硬件", "演示可在本地 Docker 环境运行，也可在 1 台 8C16G 云服务器上模拟 2 个语义节点。"],
                ["软件", "Docker Compose、Vue 3 + Vite、Spring Boot、MySQL、Go Agent、Prometheus、Grafana、SkyWalking。"],
                ["网络", "容器网络互通；演示环境需要能够访问 GitHub、镜像源和 Maven/NPM/Go 代理时才便于首次构建。"],
                ["入口", "前台统一入口 15173；后端 18081；Prometheus 19090；Grafana 13000；SkyWalking UI 18082。"],
                ["浏览器与展示", "工作台需要正常访问本地代理路径 `/api`、`/prometheus`、`/grafana`，并允许嵌入 Grafana 与 SkyWalking 可视化页面。"],
            ],
            [1.25, 4.9],
        ),
        table(
            [
                ["关键软件或镜像", "当前版本", "仓库依据"],
                ["Node 运行时", "`node:22-alpine`", "`docker-compose.yml`。"],
                ["Vue", "3.5.13", "`frontend/package.json`。"],
                ["Vite", "8.0.14", "`frontend/package.json`。"],
                ["Spring Boot", "3.3.5", "`backend/pom.xml`。"],
                ["Java", "21", "`backend/pom.xml` 与镜像配置。"],
                ["MySQL", "8.4", "`docker-compose.yml`。"],
                ["Go", "1.23.0", "`agent/go.mod`。"],
                ["Prometheus", "v2.54.1", "`docker-compose.yml`。"],
                ["Grafana", "11.2.2", "`docker-compose.yml`。"],
                ["SkyWalking OAP / UI", "10.0.1", "`docker-compose.yml`。"],
            ],
            [1.55, 1.15, 3.2],
        ),
        image("pics/总览.png", 6.2, "总览页示意"),
    ],
    "2.3 设计和实现上的限制": [
        para("服务识别基于静态规则。它适合当前演示链路，不等于对任意程序都能自动建模。"),
        para("Prometheus 的 targets 是静态配置，不会因为 Agent 发现了新服务就自动增加抓取项。"),
        para("当前环境重点是可演示、可复现，不追求生产级伸缩、审计和多租户隔离。"),
        para("系统能发现某服务，并不自动代表它已经具备完整指标接入。演示环境中的 MySQL 被刻意保留为未接独立 exporter 的状态，用于验证异常识别。"),
    ],
    "2.4 假设和约束（依赖）": [
        para("假设容器环境可用，系统时间基本同步，后端数据库可连通，SkyWalking OAP 与 UI 正常启动。"),
        para("链路展示还依赖业务流量。若没有 `demo-chain` 请求，Trace 页面只能看到组件存在，不能看到完整业务闭环。"),
        para("自动部署能力还依赖目标服务器预置环境和密钥配置，这属于部署阶段的外部约束，不是单靠应用代码就能满足的。"),
    ],
    "3、功能性需求分析": [
        para("功能性需求可以分成五组：纳管、采集、展示、下钻和验收。前四组决定系统是否像一个平台，第五组决定它是否具备可交付性。"),
        image("pics/腾讯文档-心跳上报流程图.png", 5.2, "Agent 心跳上报流程图"),
        para("F01 到 F03 的核心就是把注册、服务发现和心跳采集做成稳定链路。节点只有先完成注册并持续上报，后续的总览、详情和趋势页面才有可信数据。"),
        table(
            [
                ["编号", "需求", "说明", "验收方式"],
                ["F01", "节点注册", "Agent 启动后必须向后端注册节点名、主机名、IP、OS、版本和当前识别到的服务。", "调用 `POST /api/agents/register` 并在 `/api/nodes` 中看到节点。"],
                ["F02", "服务发现", "Agent 必须识别 Spring Boot、MySQL、Redis、Nginx、node_exporter 等固定类型服务，并带上端口；有抓取入口的服务带上指标入口，缺少抓取入口的服务保留异常标记。", "查看后端服务目录和前台服务页。"],
                ["F03", "心跳与资源采集", "Agent 必须定时上报状态、CPU、内存、磁盘和网络指标。", "查看 `/api/nodes/{id}` 与节点详情页。"],
                ["F04", "统一总览", "平台首页必须展示节点、服务、异常和趋势信息，并给出下钻入口。", "访问 `/overview` 页面。"],
                ["F05", "节点与服务下钻", "用户必须能从节点和服务两个入口查看基础信息、运行指标和原始监控链接。", "访问 `/nodes/:id`、`/services/:id`。"],
                ["F06", "链路追踪", "系统必须对 `demo-chain` 业务链路提供 Trace 摘要，并能进入 SkyWalking 原视图。", "访问 `/tracing` 和 `/api/tracing/summary`。"],
                ["F07", "调试能力", "系统必须保留 Prometheus targets 和 query 的调试入口，便于确认抓取链路是否正常。", "访问 `/debug` 和 Prometheus targets。"],
                ["F08", "部署验收", "push 到 main 后，系统应先跑测试和 smoke test，再执行远端部署脚本。", "查看 `deploy-main.yml` 与 `deploy-main.sh`。"],
            ],
            [0.7, 1.0, 2.55, 2.0],
        ),
        table(
            [
                ["服务类型", "主要识别依据", "当前示例端口", "当前指标入口状态"],
                ["SPRING_BOOT", "`java` 进程与 Spring 启动特征", "8081 / 8082", "`/actuator/prometheus`"],
                ["NGINX", "进程名含 `nginx`", "80", "`:9113/metrics`"],
                ["MYSQL", "进程名含 `mysqld` 或 `mariadbd`", "3306", "刻意不接独立 exporter"],
                ["REDIS", "进程名含 `redis-server`", "6379", "`:9121/metrics`"],
                ["NODE_EXPORTER", "进程名含 `node_exporter`", "9100", "`/metrics`"],
            ],
            [1.2, 1.95, 1.15, 2.1],
        ),
        table(
            [
                ["业务场景", "相关需求编号", "涉及页面或接口"],
                ["新节点接入平台", "F01、F02", "`/api/agents/register`、节点页、服务页。"],
                ["巡检主机健康状态", "F03、F04", "`/overview`、`/nodes`、`/nodes/{id}`。"],
                ["确认某服务有没有接全监控", "F02、F05、F07", "`/services/{id}`、Prometheus targets、调试页。"],
                ["验证业务调用链是否形成", "F06", "`/tracing`、`/api/tracing/summary`、SkyWalking UI。"],
                ["验证主分支交付闭环", "F08", "`deploy-main.yml`、`deploy-main.sh`、`tests/smoke-test.sh`。"],
            ],
            [1.4, 1.4, 3.7],
        ),
        image("pics/UML-统一工作台页面映射图.png", 5.8, "统一工作台页面映射图"),
        para("页面映射图把功能需求落到了具体入口。总览负责平台态势，节点与服务负责实体下钻，链路页负责业务闭环，调试页负责回证底层抓取事实。"),
        image("pics/服务详情（基础信息）.png", 5.8, "服务详情中的基础信息与监控入口"),
    ],
    "4、非功能需求分析": [
        para("非功能要求集中在可演示、可复现、可定位问题三点。课程验收更关心系统能否稳定展示、能否重复部署、能否快速回证问题来源。"),
        para("没有实测数据支持的吞吐量和容量指标不写入需求说明，文档只保留仓库已经体现的约束。"),
    ],
    "4.1 性能需求": [
        para("系统需支持 10 秒级心跳和 15 秒级指标抓取，满足课程演示和平台巡检需要。"),
        para("总览趋势按 5 分钟粒度保存平台快照，用来观察状态变化，不要求替代高精度时序分析。"),
        para("页面层面的性能诉求也比较明确：总览、节点和服务详情要以聚合数据快速响应，不能要求使用者手工到多个原始后台拼装信息。"),
    ],
    "4.2 可靠性需求": [
        para("心跳状态值必须经过合法性校验。缺失指标应尽量和上一条指标合并，避免图表因单次采集不完整而明显断裂。"),
        image("pics/腾讯文档-Agent心跳状态流转图.png", 5.4, "Agent 心跳状态流转图"),
        para("可靠性要求最终会落到状态解释上。心跳正常和心跳超时这两类情况需要被明确区分，否则使用者无法准确判断当前问题属于数据缺失还是节点异常。"),
        para("部署脚本需要在服务拉起后做健康检查，并确认节点注册、Prometheus targets 和业务 Trace 已经出现。"),
        para("需求层面对可靠性的判断标准也应写清：不是单看容器是否启动，而是看注册、抓取、链路和页面入口这几条关键路径能否同时成立。"),
    ],
    "4.3 安全性需求": [
        para("部署链路需区分“本机到服务器”和“服务器到 GitHub”两条 SSH 密钥路径，避免混用。"),
        para("Grafana 的匿名只读访问仅服务于演示嵌入，不应被误认为生产环境默认策略。"),
        para("调试入口和观测后台的开放程度以演示便利为先，因此文档必须明确这些配置是教学场景折中，而不是通用上线建议。"),
    ],
    "4.4 可维护性需求": [
        para("模块目录、配置文件和脚本应按职责分开。当前仓库已经按前端、后端、Agent、演示节点、基础设施、部署与测试拆分。"),
        para("常见回归点应尽量写成脚本。当前仓库已经为 SkyWalking agent 来源、Prometheus route-prefix 和部署观测配置提供了独立回归脚本。"),
        para("前端服务层、后端服务层和 Agent 模块边界保持清楚，团队成员才能分工开发而不互相覆盖。"),
    ],
    "4.5 可扩展性需求": [
        para("系统应允许继续增加服务识别规则、更多 Grafana dashboard、更多前台面板组和更多服务详情指标。"),
        para("后端对节点、服务、趋势和 Trace 摘要都用了独立 DTO，便于继续扩展读接口。"),
        para("后续继续扩展时，需求优先级应放在补更多 exporter、补异常场景、补自动化验证，而不是先追求更复杂的装饰性页面。"),
    ],
    "4.6 数据一致性": [
        para("注册接口只负责节点与服务目录，心跳接口只负责状态和资源指标，两类数据不能混在一条写路径里。"),
        para("趋势页读取的快照数据应和平台总览口径一致，避免一个页面说正常、另一个页面说异常。"),
        para("同样地，服务详情页里的抓取入口、节点详情页里的服务清单和调试页里的 targets 应当描述同一套接入事实，否则使用者无法判断是页面错、后端错还是底层抓取错。"),
    ],
    "5、验收要点": [
        para("验收证据包括代码、脚本、接口、截图和工作流运行记录。"),
        para(f"验收时可直接使用 {REPO_URL} 复核源码和工作流，用 {DEPLOY_PUBLIC_URL} 复核在线系统；当前服务器到期日期为 {DEPLOY_EXPIRES}。"),
        para("仓库对 `main` 分支启用了保护规则：禁止删除和强推，改动需通过 Pull Request 合入，至少 1 个评审通过并解决评审线程，且 `frontend-build`、`backend-test`、`agent-test` 三项状态检查必须通过。"),
        para(f"用于验收的前台页面截图以公网部署环境为准，直接取自 {DEPLOY_PUBLIC_URL} 的实际访问结果，不以本地开发环境截图代替。"),
        table(
            [
                ["验收项", "通过标准", "对应证据"],
                ["Agent 注册", "后端能收到 `app-node` 和 `middleware-node`。", "`tests/smoke-test.sh`、`deploy/deploy-main.sh`。"],
                ["服务发现", "服务目录至少包含 Spring Boot 与 MySQL。", "`tests/smoke-test.sh` 中的服务检查。"],
                ["Prometheus 抓取", "targets 中能看到 `sample-service`。", "`tests/prometheus-route-prefix-test.sh`、`tests/smoke-test.sh`。"],
                ["链路追踪", "`/api/tracing/summary` 中出现 `/api/demo-chain`。", "`deploy-main.sh` 中的 `wait_for_business_traces`。"],
                ["统一入口", "总览、节点、服务、链路、调试五类页面可访问。", "前端路由和答辩截图。"],
                ["自动部署", "main 分支工作流在 smoke test 通过后再执行 SSH 部署。", "`.github/workflows/deploy-main.yml`。"],
            ],
            [1.2, 2.3, 2.7],
        ),
        table(
            [
                ["不通过或需说明情况", "原因"],
                ["只展示静态页面，没有真实节点和业务流量", "无法证明 Agent、Prometheus 和 SkyWalking 已形成真实闭环。"],
                ["只能进入 Grafana 或 SkyWalking 原生页面，前台没有统一入口", "无法体现“一体化工作台”的项目目标。"],
                ["没有测试与部署脚本支撑", "难以证明系统具备可复现和可交付性。"],
            ],
            [2.5, 3.95],
        ),
        image("pics/GitHub actions.png", 5.8, "GitHub Actions 运行结果"),
    ],
}


DESIGN_CONTENT = {
    "1、引言": [
        para("设计范围覆盖当前仓库已经落地的总体结构、核心数据流、接口分工和数据库设计。未落地设想不写入设计文档。"),
        para("结构分层、数据流向、关键接口、状态解释逻辑和数据库职责都已经在仓库中有对应实现，设计说明直接围绕这些真实模块展开。"),
    ],
    "1.1 编写目的": [
        para("设计文档直接对应当前仓库的部署结构、数据流、接口分工和数据库职责。"),
        para(f"设计说明对应的源码仓库为 {REPO_URL}，当前在线部署入口为 {DEPLOY_PUBLIC_URL}，该云服务器租期到 {DEPLOY_EXPIRES}。"),
        para("设计说明覆盖的对象包括三类信息：节点基础信息、主机资源利用率、服务目录与调用链信息。文档中的结构设计和接口设计都围绕这三类信息如何采集、聚合和展示展开。"),
        para("系统设计围绕采集、聚合、查询和展示四层协作关系展开。"),
    ],
    "1.2 概要设计": [
        para("系统采用“自研 Agent + 聚合后端 + 统一前台 + 第三方观测底座”的组合方案。自研部分只做课题真正需要的平台控制面，把指标存储和链路可视化交给成熟组件。"),
        para("这种分层让职责比较清楚。Agent 只关心发现和上报，后端只关心聚合与判定，Prometheus/Grafana/SkyWalking 只关心各自擅长的数据能力，前端负责把它们组织成统一入口。"),
        para("统一工作台通过 `/api`、`/prometheus`、`/grafana` 三类入口和公开的 SkyWalking 地址把多个系统收口。浏览器侧先看平台聚合结果，再按需进入原生后台继续核对。"),
        table(
            [
                ["设计目标", "当前实现方式"],
                ["统一入口", "前端通过总览、节点、服务、链路、调试五类页面收口多个后台入口。"],
                ["真实数据链路", "Agent 注册与心跳、Prometheus 抓取、SkyWalking Trace、后端聚合接口共同形成闭环。"],
                ["可维护扩展", "后端服务层、前端服务层、Agent 模块和容器层次都有明确边界。"],
            ],
            [1.4, 5.0],
        ),
        image("pics/总览.png", 6.2, "统一工作台总览"),
    ],
    "1.3 业务流程图": [
        para("典型业务流程是：前台用户访问总览或详情页面；后端返回节点、服务、趋势或 Trace 摘要；用户继续下钻时，前台再通过嵌入或跳转方式进入 Grafana、Prometheus 或 SkyWalking。"),
        para("真实业务链路由 `sample-service` 的 `/api/demo-chain` 发起，请求进入 `middleware-service` 后继续访问 MySQL、Redis 和 Nginx。Prometheus 抓取指标，SkyWalking 采集 Trace，后端再对 Trace 做摘要。"),
        para("系统同时存在四条数据流：注册流负责主数据建立，心跳流负责状态和资源更新，快照流负责趋势沉淀，Trace 摘要流负责把原始链路结果加工成前台可直接阅读的结构。"),
        para("部署脚本和演示节点还补了一条预热流。`traffic-generator` 与 `deploy-main.sh` 会持续触发或主动 warmup `demo-chain`，让 Trace 页面更容易稳定看到真实业务链路。"),
        image("pics/腾讯文档-心跳上报流程图.png", 5.2, "Agent 心跳上报流程图"),
        para("这条流程把发现、上报、聚合和展示串在一起。Agent 先采集节点与服务信息，再定时上报状态和资源指标，后端完成归一化处理后由前台页面和趋势接口消费。"),
        image("pics/UML-Agent注册与心跳时序图.png", 5.8, "Agent 注册与心跳时序图"),
        para("时序图补充了接口层细节。注册负责建立节点与服务目录，心跳负责刷新状态与资源指标，两个接口职责分离，减少了写路径混杂。"),
        image("pics/架构图（观测数据流）.png", 5.8, "观测数据流"),
    ],
    "2、系统总体结构设计": [
        para("从部署拓扑看，系统分为平台控制面、观测底座和两个语义节点。平台控制面负责统一入口和平台元数据；观测底座负责指标与链路；语义节点负责生成真实业务和被观测对象。"),
        para("这种结构的关键不是组件数量，而是交互关系清楚：Agent 面向后端上报，Prometheus 面向指标入口抓取，SkyWalking 面向 Java Agent 接收 Trace，前端同时消费后端聚合数据和观测系统可视化入口。"),
        para("统一工作台在结构上承担的是“排障入口编排器”角色。总览、节点、服务、链路、调试五类页面不是并列孤岛，而是按先看全局、再看对象、再看链路、最后回证抓取的顺序组织。"),
        image("pics/腾讯文档-总体架构图.png", 5.6, "平台总体架构图"),
        para("总体架构图强调部署关系：平台控制面居中，观测底座提供共用能力，两台语义节点既产生业务流量，也承担被观测对象。"),
        image("pics/腾讯文档-服务发现与识别模块图.png", 5.6, "服务发现与识别模块图"),
        para("服务发现与识别模块对应 Agent 的 `discovery` 逻辑。当前规则主要依据进程名、监听端口和已知指标路径推断服务类型，再把结果随注册数据写入后端。"),
        table(
            [
                ["服务类型", "主要识别依据", "当前示例端口", "当前指标入口状态"],
                ["SPRING_BOOT", "`java` 进程与 Spring 启动特征", "8081 / 8082", "`/actuator/prometheus`"],
                ["NGINX", "进程名含 `nginx`", "80", "`:9113/metrics`"],
                ["MYSQL", "进程名含 `mysqld` 或 `mariadbd`", "3306", "刻意不接独立 exporter"],
                ["REDIS", "进程名含 `redis-server`", "6379", "`:9121/metrics`"],
                ["NODE_EXPORTER", "进程名含 `node_exporter`", "9100", "`/metrics`"],
            ],
            [1.2, 1.95, 1.15, 2.1],
        ),
        image("pics/架构图（角色与语义拓扑）.png", 5.8, "角色与语义拓扑"),
        table(
            [
                ["子系统", "组成", "职责"],
                ["平台控制面", "frontend、backend、mysql", "提供统一入口、聚合 API 和平台元数据持久化。"],
                ["观测底座", "prometheus、grafana、skywalking-oap、skywalking-ui", "负责指标抓取、仪表盘展示、Trace 存储与查询。"],
                ["app-node", "sample-service、node_exporter、monitor-agent、traffic-generator、SkyWalking Java Agent", "承担业务入口和持续造流量。"],
                ["middleware-node", "middleware-service、mysql、redis、nginx、exporters、monitor-agent、SkyWalking Java Agent", "承担依赖服务和被观测中间件。"],
            ],
            [1.2, 2.5, 2.4],
        ),
        table(
            [
                ["组件", "部署位置", "主要接口或暴露方式", "被谁消费"],
                ["monitor-agent", "app-node / middleware-node", "`/api/agents/register`、`/api/agents/heartbeat`", "backend。"],
                ["sample-service", "app-node", "`/actuator/prometheus`、`/api/demo-chain`", "Prometheus、业务用户、SkyWalking。"],
                ["middleware-service", "middleware-node", "`/actuator/prometheus`、`/api/middleware/profile`", "sample-service、Prometheus、SkyWalking。"],
                ["Prometheus", "observability profile", "`/api/v1/targets`、PromQL 查询", "Grafana、前台调试页、后端辅助定位。"],
                ["SkyWalking OAP / UI", "observability profile", "GraphQL、原生链路可视化", "后端 Trace 摘要、前台链路页。"],
            ],
            [1.3, 1.2, 2.1, 1.9],
        ),
        table(
            [
                ["工作台页面", "主要数据接口", "主要外部入口", "设计意图"],
                ["总览", "`/api/overview`、`/api/trends`", "Grafana / Prometheus / SkyWalking", "先看平台摘要，再按异常项下钻。"],
                ["节点页 / 节点详情", "`/api/nodes`、`/api/nodes/{id}`", "Grafana node-detail、Prometheus `up{job=...}`、SkyWalking", "把主机状态、资源和服务清单收拢到一个入口。"],
                ["服务页 / 服务详情", "`/api/services`、`/api/services/{id}`", "Grafana service-detail、Prometheus 原始查询、SkyWalking", "把服务基础信息、运行指标和调用链打通。"],
                ["链路页", "`/api/tracing/summary`", "SkyWalking 原生工作区", "先看结构化摘要，再进入 span 级细节。"],
                ["调试页", "Prometheus API 摘要", "Prometheus targets 与 graph", "用于确认指标源、标签和抓取链路。"],
            ],
            [1.25, 1.7, 1.7, 2.15],
        ),
        image("pics/UML-统一工作台页面映射图.png", 5.8, "统一工作台页面映射图"),
        para("页面映射图说明统一入口不是简单罗列菜单，而是按巡检、下钻、追踪和回证的顺序组织。总览、节点、服务、链路、调试五类页面分别对应不同排障层次。"),
        image("pics/节点详情（资源监控）.png", 6.2, "节点详情中的资源监控与下钻入口"),
        para("Compose profile 的拆分也是结构设计的一部分。核心开发层用于前后端和数据库独立联调，observability 与 nodes 两个 profile 用于补齐完整观测闭环。"),
        para("当前部署形态保持演示环境取向。`backend` 通过 `spring-boot:run` 启动，`frontend` 通过 `vite dev` 启动，SkyWalking OAP 使用 H2 存储。"),
        para("Agent 的采集方式也限定了当前结构边界。CPU、内存、网络和磁盘数据直接读取 Linux `/proc` 与系统命令，因此被观测节点当前默认假定为 Linux，而不是跨平台通用 Agent。"),
    ],
    "3、类图及接口设计": [
        para("类图和接口设计重点说明真正承担平台职责的模块和接口，不展开全部类的 UML。"),
        para("后端、前端和 Agent 三端都存在清晰的“输入 - 加工 - 输出”结构。后端负责聚合解释，前端负责工作台组织，Agent 负责主机侧发现和采样；三端边界分明，能支持组内分工并行推进。"),
        table(
            [
                ["核心类/模块", "职责", "关键字段或方法", "关联关系"],
                ["ManagedNode", "保存节点身份与当前状态。", "nodeName、hostname、ipAddress、status、lastHeartbeatAt", "一对多关联 DiscoveredService。"],
                ["DiscoveredService", "保存某节点上的服务目录。", "serviceName、serviceType、port、metricsPath、metricsPort", "多对一关联 ManagedNode。"],
                ["NodeMetrics", "保存节点级资源时序。", "cpuUsage、memoryUsage、diskUsage、networkRxMbps、networkTxMbps", "多对一关联 ManagedNode。"],
                ["HeartbeatEvent", "记录节点心跳历史。", "status、createdAt", "多对一关联 ManagedNode。"],
                ["MetricsSnapshot", "保存平台级趋势快照。", "totalNodes、warningNodes、healthyServices、unresolvedAlerts", "独立表，用于趋势图。"],
                ["NodeRegistryService", "处理注册、心跳、聚合、异常和外链。", "registerNode、heartbeat、overview、getNode、getService、getTrends", "调用各 Repository，向 PortalController 与 AgentController 提供服务。"],
                ["SkyWalkingQueryService", "查询 SkyWalking GraphQL 并生成 Trace 摘要。", "loadTracingSummary、loadLatestTracePreview", "被 PortalController 调用。"],
            ],
            [1.2, 1.6, 1.85, 1.55],
        ),
        image("pics/UML-后端核心实体类图.png", 5.8, "后端核心实体类图"),
        para("类图把节点、服务、心跳、资源时序和平台快照的关系压缩到一张图里。这里的主干是 ManagedNode，其他实体围绕节点主数据和平台快照展开。"),
        table(
            [
                ["Agent 模块", "职责", "输入", "输出"],
                ["config", "读取环境变量与运行配置。", "容器环境变量。", "Agent 运行参数。"],
                ["system", "采集主机名、IP、OS。", "操作系统信息。", "节点基础信息。"],
                ["discovery", "扫描进程与端口并识别服务。", "`ps`、`ss -ltnp`。", "服务目录与抓取入口信息。"],
                ["metrics", "采集 CPU、内存、磁盘、网络。", "`/proc`、`df`。", "心跳中的资源指标。"],
                ["httpclient", "调用后端注册与心跳接口。", "JSON payload。", "平台注册与状态刷新结果。"],
            ],
            [1.25, 1.8, 1.5, 1.95],
        ),
        table(
            [
                ["前端模块", "职责", "主要消费的数据或入口"],
                ["`api.js`", "统一调用后端 JSON API。", "总览、节点、服务、趋势和 Trace 摘要接口。"],
                ["`prometheus.js`", "调用 Prometheus API 并生成 targets 摘要。", "调试页与抓取健康度展示。"],
                ["`observability.js`", "统一构造 Grafana、Prometheus、SkyWalking 跳转与嵌入 URL。", "总览页、详情页、链路页、调试页。"],
                ["页面层 `pages/`", "按排障视角组织工作台。", "`/overview`、`/nodes`、`/services`、`/tracing`、`/debug`。"],
            ],
            [1.45, 2.2, 2.85],
        ),
        table(
            [
                ["接口", "方法", "作用", "主要输入输出"],
                ["/api/agents/register", "POST", "节点注册与服务同步。", "输入节点基础信息和服务目录；输出节点详情。"],
                ["/api/agents/heartbeat", "POST", "状态与资源指标上报。", "输入状态与主机资源数据；输出节点摘要。"],
                ["/api/overview", "GET", "总览页聚合数据。", "输出节点计数、服务计数、异常和告警摘要。"],
                ["/api/nodes", "GET", "节点列表。", "支持按状态、关键字、服务类型和排序方式查询。"],
                ["/api/nodes/{id}", "GET", "节点详情。", "输出节点摘要、资源指标、服务清单和快捷跳转链接。"],
                ["/api/services", "GET", "服务列表。", "输出服务摘要、所属节点、异常状态等。"],
                ["/api/services/{id}", "GET", "服务详情。", "输出运行指标入口、调用链入口、基础信息和调试链接。"],
                ["/api/trends", "GET", "趋势数据。", "按时间窗返回平台快照趋势。"],
                ["/api/tracing/summary", "GET", "Trace 摘要。", "返回最近业务 Trace、服务链和依赖链摘要。"],
                ["/api/demo-chain", "GET", "演示业务入口。", "调用下游 middleware-service 并返回聚合结果。"],
                ["/api/middleware/profile", "GET", "演示依赖访问。", "访问 MySQL、Redis、Nginx 并返回结果。"],
            ],
            [1.55, 0.7, 1.65, 2.25],
        ),
        table(
            [
                ["关键数据流", "起点", "中转处理", "终点"],
                ["注册流", "Agent 启动", "NodeRegistryService 同步节点与服务清单", "节点页、服务页。"],
                ["心跳流", "Agent 定时上报", "NodeRegistryService 归一化状态与资源指标", "总览页、节点详情页。"],
                ["快照流", "MetricsCollectionScheduler", "保存 `MetricsSnapshot`", "趋势图与平台摘要。"],
                ["Trace 摘要流", "SkyWalking OAP", "SkyWalkingQueryService 查询、过滤、合并", "链路页与服务依赖摘要。"],
            ],
            [1.15, 1.25, 2.2, 1.9],
        ),
        table(
            [
                ["规则", "判定位置", "平台表现"],
                ["注册成功后节点默认记为 ONLINE", "`registerNode`", "节点进入纳管目录并可立即出现在节点页。"],
                ["ONLINE 但还没有心跳时间", "`effectiveStatus`、`buildStatusSummary`", "节点显示等待心跳或 warning 提示。"],
                ["最近心跳超过 90 秒", "`isHeartbeatTimeout`、`effectiveStatus`", "节点转入 WARNING，并提示心跳超时风险。"],
                ["服务缺少 metricsPath", "`isAbnormalService`", "服务被视为异常，前台标记为 `NO_SCRAPE_ENDPOINT`。"],
                ["MySQL 演示场景", "middleware-node 中的 MySQL", "刻意保留为缺少抓取入口，用于验证异常识别。"],
                ["节点详情和服务详情统一生成外链", "`buildQuickLinks`、`buildServiceQuickLinks`", "可直接跳到 Grafana、Prometheus、SkyWalking。"],
            ],
            [1.75, 1.45, 3.1],
        ),
        image("pics/UML-节点状态判定状态图.png", 5.4, "节点状态判定状态图"),
        para("状态图把 `effectiveStatus` 等规则的结果直观化。当前实现重点区分注册成功、持续心跳正常和心跳超时告警，不把所有异常都折叠成同一种状态。"),
        para("接口设计上的一个明显特点是读写分离。`/api/agents/register` 和 `/api/agents/heartbeat` 只承担写入职责，PortalController 下面的 `/api/overview`、`/api/nodes`、`/api/services`、`/api/tracing/summary` 等接口则面向前台读场景返回聚合 DTO。"),
        para("注册接口内部按 `nodeName` 定位节点。重复注册会更新该节点的基础信息和服务目录，而不是插入重复主机记录；这使节点名在当前设计中承担了注册幂等键的作用。"),
        para("列表与趋势接口也体现了当前规模假设。节点列表在服务层内存过滤并支持按最近心跳排序，趋势查询最大窗口为 720 小时，前台则提供 15 分钟到 24 小时的常用快捷范围。"),
        para("这种分工能减少接口语义混乱：Agent 不需要关心前台如何展示，前台也不需要理解主机采样细节，只要消费已经整理好的节点、服务、趋势和 Trace 摘要结果即可。"),
        para("前端的 `observability.js` 还负责第二层组织工作。它根据节点名、服务名、指标端口和服务类型拼装 Grafana、Prometheus、SkyWalking 的嵌入地址与回跳地址，使页面组件不必直接处理复杂的外部 URL 规则。"),
        image("pics/服务详情（运行指标）.png", 5.8, "服务详情中的运行指标面板"),
        image("pics/SkyWalking Topology.png", 5.8, "SkyWalking 服务拓扑"),
    ],
    "4、数据库设计": [
        para("后端通过 JPA `ddl-auto=update` 维护表结构。当前平台的数据库更像平台元数据与快照存储，而不是海量时序库。"),
        para("数据库设计体现的是平台管理视角：节点、服务、状态事件、资源时序和平台级快照分别落在不同实体里，既避免把所有信息揉成一张大表，也方便不同页面按职责读取。"),
        image("pics/腾讯文档-数据库关系图.png", 5.2, "平台数据库关系图"),
        para("数据库关系图展示了节点、服务、心跳、资源采样和平台快照之间的协作关系。节点是主干，服务、心跳和资源采样围绕节点展开，平台快照保持独立，用于平台级历史统计。"),
        para("`managed_nodes` 表保存被纳管节点的基础信息。"),
        table(
            [
                ["字段", "类型", "约束", "说明"],
                ["id", "Long", "主键", "节点主键。"],
                ["node_name", "String", "唯一、非空", "节点标识。"],
                ["hostname", "String", "非空", "主机名。"],
                ["ip_address", "String", "非空", "节点 IP。"],
                ["os_name", "String", "非空", "操作系统名。"],
                ["agent_version", "String", "非空", "Agent 版本。"],
                ["status", "String", "非空", "当前状态。"],
                ["last_seen_at", "Instant", "非空", "最近一次被后端看到的时间。"],
                ["last_heartbeat_at", "Instant", "可空", "最近心跳时间。"],
            ],
            [1.35, 0.9, 1.2, 2.55],
        ),
        para("`discovered_services` 表保存节点上的服务目录。"),
        table(
            [
                ["字段", "类型", "约束", "说明"],
                ["id", "Long", "主键", "服务主键。"],
                ["service_name", "String", "非空", "服务名称。"],
                ["service_type", "String", "非空", "服务类型，如 SPRING_BOOT、MYSQL。"],
                ["port", "Integer", "非空", "服务端口。"],
                ["process_name", "String", "非空", "进程名。"],
                ["metrics_path", "String", "可空", "指标抓取路径。"],
                ["metrics_port", "Integer", "可空", "指标端口。"],
                ["node_id", "Long", "外键、非空", "所属节点。"],
            ],
            [1.35, 0.9, 1.2, 2.55],
        ),
        para("`node_metrics` 表保存节点资源时序。"),
        table(
            [
                ["字段", "类型", "约束", "说明"],
                ["id", "Long", "主键", "记录主键。"],
                ["node_id", "Long", "外键、非空", "所属节点。"],
                ["collected_at", "Instant", "非空", "采集时间。"],
                ["cpu_usage", "Double", "可空", "CPU 使用率。"],
                ["memory_usage", "Double", "可空", "内存使用率。"],
                ["memory_total_mb", "Long", "可空", "内存总量。"],
                ["memory_used_mb", "Long", "可空", "内存已用。"],
                ["disk_usage", "Double", "可空", "磁盘使用率。"],
                ["disk_total_gb", "Long", "可空", "磁盘总量。"],
                ["disk_used_gb", "Long", "可空", "磁盘已用。"],
                ["network_rx_mbps / network_tx_mbps", "Double", "可空", "网络收发速率。"],
            ],
            [1.55, 0.9, 1.15, 2.4],
        ),
        para("`heartbeat_events` 表保存心跳历史，`metrics_snapshots` 表保存平台级趋势快照。"),
        table(
            [
                ["表", "关键字段", "作用"],
                ["heartbeat_events", "node_id、status、created_at", "保留节点状态历史，便于后续审计和分析。"],
                ["metrics_snapshots", "timestamp、total_nodes、warning_nodes、healthy_services、unresolved_alerts", "支撑总览页趋势图。"],
            ],
            [1.4, 2.15, 2.65],
        ),
        para("`metrics_snapshots` 还对 `timestamp` 建立了正序和倒序索引，说明该表的主要读取方式就是按时间窗口取最近一段快照，并支持清理旧数据时快速定位时间边界。"),
        table(
            [
                ["关系", "说明"],
                ["ManagedNode -> DiscoveredService", "一个节点可对应多个已发现服务，服务从节点视角被纳管。"],
                ["ManagedNode -> HeartbeatEvent", "一个节点会持续产生多条状态事件，用于保留时序历史。"],
                ["ManagedNode -> NodeMetrics", "一个节点会持续产生多条资源采样记录，用于详情页资源趋势。"],
                ["MetricsSnapshot 独立存在", "它不是某个节点的附属表，而是平台整体在某一时刻的管理快照。"],
            ],
            [2.2, 4.3],
        ),
        table(
            [
                ["数据对象", "写入时机", "默认保留策略", "主要用途"],
                ["`node_metrics`", "Agent 心跳携带指标时写入", "7 天，定时清理", "节点详情页资源趋势与最新资源摘要。"],
                ["`metrics_snapshots`", "后端定时任务每 5 分钟写入", "30 天，定时清理", "总览趋势图和平台级历史统计。"],
                ["`heartbeat_events`", "每次心跳写入", "随数据库保留", "保留节点状态历史。"],
                ["`discovered_services`", "每次注册时同步", "随节点主数据维护", "服务目录和服务详情基础信息。"],
            ],
            [1.5, 1.8, 1.5, 1.8],
        ),
        para("如果从读侧用途来理解这几张表，`managed_nodes` 和 `discovered_services` 决定平台静态目录，`heartbeat_events` 和 `node_metrics` 决定节点动态状态，`metrics_snapshots` 则负责给总览页提供平台级历史视角。"),
        para("定时任务还负责补齐历史节点缺失的心跳时间，以及按周期清理旧快照和旧节点指标。数据库层因此不仅保存数据，还承担平台管理口径的一部分维护职责。"),
        para("这种分层存储也说明当前数据库承担的是“平台管理数据”而不是“完整可观测时序库”。Prometheus 和 SkyWalking 继续负责底层大规模指标与链路存储，平台后端只保留自己需要的管理摘要。"),
    ],
    "5、其他": [
        para("部署设计上，仓库使用 Compose profile 把核心开发层、观测层和演示节点层拆开。这样可以在联调时只起必要组件，在答辩或验收时再补起完整观测环境。"),
        para(f"当前外部交付形态也比较明确：源码、工作流和文档统一放在 {REPO_URL}；在线系统通过 {DEPLOY_PUBLIC_URL} 提供公网访问；该云服务器当前到期日期为 {DEPLOY_EXPIRES}。"),
        para("当前设计的取舍也要保留说明：它优先解决课题闭环和展示路径，因此把更多 exporter、生产化启动参数、多节点高可用和浏览器级 E2E 测试留作后续扩展点。"),
        para("平台把 Prometheus、Grafana 和 SkyWalking 作为底层能力整合进统一工作台。额外成本主要在路径代理、嵌入与统一入口一致性。"),
        para("当前设计路线直接对应课题要求：平台控制面、统一工作台和工程闭环由自研部分承担，成熟观测组件承担指标和链路支撑。"),
    ],
}


TEST_CONTENT = {
    "1、引言": [
        para("测试证据以 GitHub Actions 工作流、仓库脚本和现有截图材料为主。没有工作流记录、脚本断言或界面证据支撑的内容，不写成“已通过”。"),
        para("测试结论分别来自自动化记录、脚本断言和界面截图。三类证据共同说明系统验证范围和当前状态。"),
    ],
    "1.1 编写目的": [
        para("课程验收同时关注页面可见和链路验证，测试文档分别给出对应证据。"),
        para("自动化、冒烟验证、部署后自检和截图回证集中放在同一份文档里，便于直接核对系统状态。"),
    ],
    "1.2 测试目标": [
        para("测试目标分三层：一是验证 Agent、后端、前端和观测组件的集成闭环；二是验证部署脚本和关键配置不会因为修改而失效；三是为答辩时的人工演示提供可回证的截图与入口。"),
        para("测试目标不是单看某个接口是否返回 200，而是验证节点能否接入、指标能否抓到、链路能否形成、工作台能否展示、部署能否重复。"),
    ],
    "1.3 测试环境": [
        para("测试环境并不等于“有一台机器能启动就行”。对于本项目，环境要同时满足容器编排、指标抓取、链路追踪、前台代理和自动部署几个方面的前提。"),
        table(
            [
                ["项", "内容"],
                ["代码环境", "当前仓库主分支代码。"],
                ["仓库地址", REPO_URL],
                ["运行方式", "Docker Compose；完整场景使用 `observability` 与 `nodes` profile。"],
                ["核心服务", "frontend、backend、mysql、prometheus、grafana、skywalking-oap、skywalking-ui、app-node、middleware-node。"],
                ["外部依赖", "GitHub Actions、SSH 部署、镜像源与包管理代理配置。"],
                ["公网部署入口", f"{DEPLOY_PUBLIC_URL}（服务器到期日期：{DEPLOY_EXPIRES}）"],
                ["主验证来源", "GitHub Actions 的 `ci.yml`、`deploy-main.yml` 与现有答辩截图材料。"],
                ["前台截图来源", f"前台页面截图取自公网部署地址 {DEPLOY_PUBLIC_URL} 的实际访问结果，不使用本地开发环境截图。"],
                ["浏览器与截图材料", "前台工作台截图、Prometheus 后台截图、SkyWalking Trace 截图、GitHub Actions 运行截图。"],
            ],
            [1.4, 4.8],
        ),
    ],
    "1.4 测试人员": [
        para("测试分工与开发分工保持一致。组长负责跨模块链路确认，其余成员分别覆盖各自主负责页面与功能。"),
        table(
            [
                ["人员", "主要测试关注点"],
                ["陈梦泽", "整体联调、部署链路、CI/CD、Agent/后端主线和最终演示材料。"],
                ["陈紫暄", "总览页、趋势图、节点页筛选与节点详情交互。"],
                ["曹丹", "服务视图可用性、趋势表达细节、答辩视频整理。"],
            ],
            [1.35, 4.85],
        ),
    ],
    "1.5 参考材料": [
        para("测试材料包括代码级脚本、工作流级自动化和展示级截图。三类材料共同构成测试证据。"),
        table(
            [
                ["材料", "用途"],
                [REPO_URL, "复核源码、工作流和文档版本。"],
                ["`tests/smoke-test.sh`", "完整环境冒烟验证。"],
                ["`tests/skywalking-agent-source-test.sh`", "检查 SkyWalking Java Agent 来源是否稳定。"],
                ["`tests/prometheus-route-prefix-test.sh`", "检查 Prometheus route-prefix 和 smoke test 配置。"],
                ["`tests/deploy-observability-config-test.sh`", "检查部署脚本与观测配置关键路径。"],
                ["`tests/network-mirror-config-test.sh`", "检查镜像源与代理配置，主要用于环境准备阶段复核。"],
                ["`.github/workflows/ci.yml`、`deploy-main.yml`", "查看持续集成与自动部署流程。"],
            ],
            [2.4, 3.75],
        ),
    ],
    "2、测试用例设计": [
        para("测试设计先覆盖最容易回归的配置点，再用冒烟脚本串起主链路，最后用截图和页面入口补充人工复核。"),
        table(
            [
                ["测试层级", "执行位置", "主要覆盖对象", "当前证据"],
                ["模块级测试", "GitHub Actions", "前端、后端、Agent 各自的单元或模块逻辑", "`ci.yml`、`deploy-main.yml`。"],
                ["配置回归测试", "GitHub Actions；网络镜像脚本可用于环境准备复核", "SkyWalking agent 来源、Prometheus route-prefix、部署观测配置、镜像源配置", "`tests/*.sh`。"],
                ["全链路冒烟测试", "GitHub Actions", "容器拉起、节点注册、服务发现、指标抓取、Trace 形成", "`tests/smoke-test.sh`。"],
                ["人工页面复核", "演示环境", "总览、节点、服务、链路、调试和第三方后台截图", "`pics/`、答辩 PPT。"],
            ],
            [1.15, 1.2, 2.0, 2.15],
        ),
        para("这种测试层次覆盖了课程交付最容易失分的回归点：依赖、路径、targets、业务 Trace 和部署顺序。"),
    ],
    "2.1测试数据准备": [
        para("运行完整环境前，需要准备 `.env`，并保证 Docker、镜像源和 Maven/NPM/Go 代理配置可用。"),
        para("为了让 Trace 页面稳定出数，系统在 `app-node` 内置 `traffic-generator.sh`，部署脚本也会主动访问 `/api/demo-chain` 进行预热。"),
        para("因此，本项目的测试准备不仅包括代码和依赖，还包括业务流量准备。没有持续流量时，SkyWalking 和部分 Grafana 面板会出现“组件可用但链路为空”的情况。"),
    ],
    "2.2测试策略": [
        para("当前测试策略分为四层：模块级单测、配置回归、完整环境冒烟、人工页面复核。"),
        para("其中，模块级单测由 GitHub Actions 执行；`packaging-test` 负责执行关键配置回归脚本；完整环境冒烟由 `smoke-test.sh` 验证节点注册、Prometheus 抓取和业务 Trace；人工复核主要依赖前台页面、原生后台截图和工作流运行记录。"),
        para("四层策略同时覆盖代码逻辑回归、工程配置回归、完整环境闭环和答辩展示材料。"),
    ],
    "2.3测试原则": [
        para("测试内容尽量直接对应课题要求：主机纳管、服务发现、资源监控、技术栈信息、指标抓取和调用链展示。"),
        para("对没有当前 CI 记录、截图或仓库证据支持的部分，不把它表述成“已通过”；对已经由 GitHub Actions 跑通的链路，则直接按通过状态记录。"),
        para("同样需要区分“脚本存在”和“结果已通过”。对于课程文档，前者可以证明工程意识，后者才能证明当前版本状态。两者都应写，但不能混写。"),
    ],
    "2.4测试用例设计": [
        para("测试用例设计围绕题目原始要求展开，不额外发明脱离课题的花哨场景。只要纳管、发现、采集、展示、追踪和部署验收六条主线能闭环，课程目标就能被直接回证。"),
        table(
            [
                ["用例ID", "测试目标", "执行方式", "预期结果"],
                ["T01", "节点注册", "运行完整环境后访问 `/api/nodes`。", "能看到 `app-node` 和 `middleware-node`。"],
                ["T02", "服务发现", "访问 `/api/services`。", "能看到 Spring Boot、MySQL、Redis、Nginx、node_exporter 等服务目录。"],
                ["T03", "心跳与资源指标", "访问节点详情接口或前台节点详情页。", "能看到 CPU、内存、磁盘、网络等指标。"],
                ["T04", "总览页聚合", "访问 `/overview`。", "能看到节点计数、服务计数、异常入口和趋势图。"],
                ["T05", "服务详情下钻", "进入 `/services/:id`。", "能看到基础信息、运行指标、Trace 入口和 Prometheus 调试入口。"],
                ["T06", "Prometheus 抓取链路", "查看 Prometheus targets。", "能看到 `sample-service` 等 targets 正常暴露。"],
                ["T07", "Trace 采集链路", "访问 `/api/tracing/summary` 或前台链路页。", "能出现 `/api/demo-chain` 相关业务 Trace。"],
                ["T08", "部署脚本自检", "执行 `deploy/deploy-main.sh`。", "脚本会等待健康检查、节点注册、Prometheus targets 和业务 Trace。"],
                ["T09", "关键配置回归", "执行工作流中的关键 shell 脚本。", "SkyWalking agent 来源、Prometheus route-prefix 与部署观测关键路径保持一致。"],
                ["T10", "服务视图分组展示", "访问 `/services`。", "服务按类型分组，缺失抓取入口的服务被单独标识。"],
                ["T11", "调试页原始链路验证", "访问 `/debug`。", "可看到 Prometheus 抓取摘要与原始调试入口。"],
                ["T12", "自动部署串行约束", "查看 `deploy-main.yml`。", "仅当 smoke test 通过后才执行 SSH 部署。"],
            ],
            [0.75, 1.35, 1.95, 2.15],
        ),
        para("这些用例同时覆盖接口、页面、工作流和部署脚本，因为本项目的交付对象不是单一应用，而是一套需要完整环境配合的系统。"),
    ],
    "3、测试过程": [
        para("GitHub Actions 已串行执行前端测试与构建、后端测试、Agent 测试、关键配置回归、完整 smoke test 和 main 分支自动部署，现有记录显示这些流程通过。"),
        para("主分支工作流中的 `packaging-test` 执行了 `skywalking-agent-source-test.sh`、`prometheus-route-prefix-test.sh` 和 `deploy-observability-config-test.sh`，用于拦截最容易影响演示闭环的配置回归。"),
        para("仓库对 `main` 分支启用了保护规则：禁止删除和强推，改动需通过 Pull Request 合入，至少 1 个评审通过并解决评审线程；`frontend-build`、`backend-test`、`agent-test` 三项状态检查必须通过，并要求按最新主分支严格复核。"),
        para(f"除了工作流记录，当前还保留了在线复核入口：公网部署地址为 {DEPLOY_PUBLIC_URL}，该云服务器到期日期为 {DEPLOY_EXPIRES}。在到期日前，可直接通过该地址复核当前部署版本。"),
        para("`smoke-test.sh` 负责拉起完整环境并检查节点注册、服务发现、Prometheus targets 和业务 Trace；`deploy-main.sh` 则在 smoke test 通过后继续执行远端部署、健康检查和业务链路预热。"),
        image("pics/UML-主分支交付时序图.png", 5.8, "主分支交付与部署时序图"),
        image("pics/GitHub action流程详情.png", 5.9, "GitHub Actions 中的测试与部署依赖关系"),
        para("这条时序和当前交付链路的依赖关系一致：先完成基础测试，再做配置回归和完整冒烟，最后才进入远端部署。这样可以避免把明显有问题的版本直接推到服务器。"),
        para("GitHub Actions 流程详情图进一步把这一点固定成了可见事实：`frontend-build`、`backend-test`、`agent-test` 和 `packaging-test` 全部成功后，工作流才会放行 `smoke-test`；`smoke-test` 通过后，才会进入 `deploy`。"),
        table(
            [
                ["工作流 / 作业", "作用", "当前状态口径"],
                ["`ci.yml` / `frontend-build`", "执行前端依赖安装、测试与构建。", "GitHub Actions 已通过。"],
                ["`ci.yml` / `backend-test`", "执行后端 Maven 测试。", "GitHub Actions 已通过。"],
                ["`ci.yml` / `agent-test`", "执行 Agent Go 测试。", "GitHub Actions 已通过。"],
                ["`deploy-main.yml` / `packaging-test`", "执行关键配置回归脚本。", "GitHub Actions 已通过。"],
                ["`deploy-main.yml` / `smoke-test`", "验证完整环境闭环。", "GitHub Actions 已通过。"],
                ["`deploy-main.yml` / `deploy`", "在 smoke test 通过后执行远端部署。", "GitHub Actions 已通过。"],
            ],
            [1.75, 2.4, 2.35],
        ),
        table(
            [
                ["脚本或检查点", "GitHub Actions 中的验证内容", "结果口径"],
                ["`tests/skywalking-agent-source-test.sh`", "检查 app-node 与 middleware-node Dockerfile 是否继续从 Apache 归档包复制 SkyWalking Java Agent。", "`packaging-test` 已通过。"],
                ["`tests/prometheus-route-prefix-test.sh`", "检查 Prometheus route-prefix 与 smoke test 健康检查路径是否一致。", "`packaging-test` 已通过。"],
                ["`tests/deploy-observability-config-test.sh`", "检查部署脚本中的 Prometheus、Grafana、SkyWalking 关键等待逻辑与业务 Trace 预热路径。", "`packaging-test` 已通过。"],
                ["`tests/smoke-test.sh`", "启动完整环境并检查节点注册、服务发现、Prometheus 抓取与业务 Trace。", "`smoke-test` 已通过。"],
                ["`deploy/deploy-main.sh`", "在 smoke test 之后执行远端部署，并等待健康检查、targets 与业务 Trace。", "`deploy` 已通过。"],
            ],
            [2.35, 2.55, 1.3],
        ),
        table(
            [
                ["关键检查点", "断言内容", "对应意义"],
                ["节点注册", "`/api/nodes` 中出现 `app-node` 和 `middleware-node`。", "证明 Agent 注册链路成立。"],
                ["服务发现", "`/api/services` 中出现 `SPRING_BOOT` 与 `MYSQL`。", "证明节点目录与服务目录已同步。"],
                ["指标抓取", "Prometheus targets 中出现 `sample-service`。", "证明 Prometheus 已接到业务指标源。"],
                ["链路形成", "`/api/tracing/summary` 中出现 `/api/demo-chain`。", "证明业务 Trace 已被 SkyWalking 和后端摘要链路接住。"],
                ["部署后预热", "`deploy-main.sh` 主动 warmup `/api/demo-chain` 并重试等待业务 Trace。", "证明远端部署不是只拉起容器，而是继续验证观测闭环。"],
            ],
            [1.4, 2.45, 2.1],
        ),
        para("工作流顺序为：先在 PR 阶段完成前端、后端、Agent 的基础回归，再在 main 工作流中补齐配置检查、完整冒烟和远端部署。"),
        image("pics/Prometheus后台（targets）.png", 5.8, "Prometheus targets 截图"),
        image("pics/Prometheus后台（service discovery）.png", 6.2, "Prometheus service discovery 截图"),
        image("pics/SkyWalking Trace.png", 5.8, "SkyWalking Trace 截图"),
        image("pics/GitHub actions.png", 5.8, "GitHub Actions 截图"),
        image("pics/调试面板.png", 5.8, "调试面板截图"),
        para("这些截图在测试文档中的作用并不相同。Prometheus targets 和 service discovery 用来证明抓取链路已经建立；SkyWalking Trace 用来证明业务链路确实形成；GitHub Actions 用来证明自动化链路已经跑通；调试面板截图则说明前台已经把原始观测入口组织成可直接排查的工作台。"),
        para(f"测试文档、报告和答辩 PPT 中使用的前台页面截图，均来自公网部署地址 {DEPLOY_PUBLIC_URL} 的实际访问结果，不使用本地开发环境页面截图。这样页面证据与在线复核入口保持一致。"),
        para("演示视频按总览、节点或服务下钻、链路追踪、原生后台回证的顺序组织，只作为展示路径补充，不替代代码、脚本和工作流证据。"),
    ],
    "4、结论": [
        para("从仓库证据看，这套项目已经具备课程设计所需的主链路：节点纳管、服务发现、主机资源采集、统一总览、服务下钻、调用链展示和自动部署验收。"),
        para("同时也要把边界说清楚：演示环境中的 MySQL 被刻意保留为未接独立 exporter 的状态，用于验证异常识别；Prometheus 仍是静态 targets；当前部署形态偏演示环境。这些边界不影响课程要求中的监控闭环和自动化验收。"),
        para("当前测试证据包括模块级测试、配置回归、冒烟验证和远端部署四层记录。"),
        para("结论口径以 GitHub Actions 运行记录、脚本断言和截图材料为准。课程设计所要求的关键链路已经被多层证据覆盖，当前版本状态可以由工作流记录和界面材料共同证明。"),
    ],
    "5、其他": [
        para("后续若要补强测试体系，优先建议增加两类内容：一类是浏览器级端到端用例，用来覆盖工作台主要路径；另一类是更多服务发现与异常场景的 Agent 单测，用来降低规则改动带来的回归风险。"),
        table(
            [
                ["补强方向", "价值"],
                ["浏览器级 E2E", "把总览、节点、服务、链路和调试五类页面串成真实用户路径。"],
                ["更多 Agent 规则单测", "降低服务发现规则调整后误识别或漏识别的风险。"],
                ["异常注入场景", "更有针对性地验证 `NO_SCRAPE_ENDPOINT`、节点告警和链路缺失等场景的页面提示。"],
            ],
            [2.1, 4.4],
        ),
        para("这几项补强都能在现有仓库结构上继续推进，不需要推翻当前平台。现有测试体系已经足够支撑课程验收，同时也给后续继续深化留下了清晰入口。"),
    ],
}


REQ_HEADING_RENAMES = {
    "1.1 项目背景": "1.1 编写目的",
    "1.2 项目目的": "1.2 项目背景",
    "1.3 项目特色": "1.3 参考资料",
    "1.4 项目风险": "1.4 术语说明",
    "1.5 产品范围": "1.5 产品范围",
    "5、项目进度安排": "5、验收要点",
}


def set_cover(document: Document, doc_type: str) -> None:
    topic = TITLE if doc_type == "报告书" else f"{TITLE}—{doc_type}"
    set_cover_line(document.paragraphs[3], COVER_TITLE)
    set_cover_line(document.paragraphs[6], topic)
    set_cover_field(document.paragraphs[10], "学  院   ", COLLEGE, 21)
    set_cover_field(document.paragraphs[11], "专  业   ", MAJOR, 22)
    set_cover_field(document.paragraphs[12], "组长姓名 ", LEADER, 21)
    set_cover_field(document.paragraphs[13], "组  员   ", MEMBERS, 19)
    set_cover_field(document.paragraphs[14], "指导教师 ", MENTOR, 22)
    set_cover_field(document.paragraphs[15], "课程编号 ", COURSE_CODE, 22)
    set_cover_field(document.paragraphs[16], "课程学分 ", CREDITS, 22)
    set_cover_field(document.paragraphs[17], "起始日期 ", START_DATE, 21)


def rename_headings(document: Document, mapping: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in mapping:
            set_paragraph_text_like_existing(paragraph, mapping[text])


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"未找到段落：{text}")


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None  # type: ignore[attr-defined]


def cleanup_blank_placeholders(document: Document, heading_texts: list[str]) -> None:
    while True:
        changed = False
        paragraphs = list(document.paragraphs)
        for i, paragraph in enumerate(paragraphs[:-1]):
            current = paragraph.text.strip()
            nxt = paragraphs[i + 1].text.strip()
            if current in heading_texts and nxt == "":
                remove_paragraph(paragraphs[i + 1])
                changed = True
                break
        if not changed:
            break


def append_paragraph_after(anchor: Paragraph, text: str, align: WD_ALIGN_PARAGRAPH | None = None, italic: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    run = paragraph.add_run(text)
    set_run_style(run, font_name=BODY_FONT, size_pt=12, italic=italic)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    if align is not None:
        paragraph.alignment = align
    return paragraph


def append_image_after(document: Document, anchor: Paragraph, image_path: Path, width_in: float, caption: str, figure_no: int) -> Paragraph:
    image_paragraph = append_paragraph_after(anchor, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    image_paragraph.paragraph_format.space_after = Pt(4)
    image_paragraph.runs[0].add_picture(str(image_path), **picture_kwargs(document, image_path, width_in))
    caption_paragraph = append_paragraph_after(
        image_paragraph,
        f"图{figure_no} {caption}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )
    if caption_paragraph.runs:
        set_run_style(caption_paragraph.runs[0], font_name=BODY_FONT, size_pt=10.5, italic=True)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    return caption_paragraph


def format_table(table_obj: Table, column_widths: list[float]) -> None:
    try:
        table_obj.style = "Table Grid"
    except KeyError:
        try:
            table_obj.style = "Normal Table"
        except KeyError:
            pass
    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_obj.autofit = False
    set_table_borders(table_obj)
    apply_column_widths(table_obj, column_widths)
    if table_obj.rows:
        mark_header_row(table_obj.rows[0])
    for row_index, row in enumerate(table_obj.rows):
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=105, bottom=105, start=135, end=135)
            if row_index == 0:
                set_cell_shading(cell, TABLE_HEADER_FILL)
            elif row_index % 2 == 0:
                set_cell_shading(cell, TABLE_ALT_FILL)
            else:
                set_cell_shading(cell, None)
            for para_obj in cell.paragraphs:
                para_obj.paragraph_format.space_after = Pt(0)
                para_obj.paragraph_format.line_spacing = 1.15
                para_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                if row_index == 0:
                    for run in para_obj.runs:
                        run.bold = True
                for run in para_obj.runs:
                    set_run_style(run, font_name=BODY_FONT, size_pt=10.5, bold=run.bold)


def append_table_after(document: Document, anchor: Paragraph, rows: list[list[str]], widths: list[float] | None = None) -> Paragraph:
    table_obj = document.add_table(rows=len(rows), cols=len(rows[0]))
    anchor._p.addnext(table_obj._tbl)
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = table_obj.cell(row_index, col_index)
            cell.text = value
    format_table(table_obj, resolve_column_widths(document, rows, widths))
    tail = OxmlElement("w:p")
    table_obj._tbl.addnext(tail)
    tail_paragraph = Paragraph(tail, anchor._parent)
    tail_paragraph.paragraph_format.space_after = Pt(6)
    return tail_paragraph


def apply_blocks(document: Document, content: dict[str, list[dict]]) -> None:
    heading_texts = list(content.keys())
    image_index = 1
    cleanup_blank_placeholders(document, heading_texts)
    for heading, blocks in content.items():
        anchor = find_paragraph(document, heading)
        for block in blocks:
            if block["type"] == "paragraph":
                anchor = append_paragraph_after(anchor, block["text"])
            elif block["type"] == "image":
                anchor = append_image_after(document, anchor, ROOT / block["path"], block["width"], block["caption"], image_index)
                image_index += 1
            elif block["type"] == "table":
                anchor = append_table_after(document, anchor, block["rows"], block.get("widths"))
            else:
                raise ValueError(f"不支持的 block 类型：{block['type']}")


def build_report_doc() -> None:
    output = REPORT_DIR / "课程设计报告书.docx"
    shutil.copyfile(BASE_DIR / "课程设计报告书.docx", output)
    document = Document(output)
    set_cover(document, "报告书")
    apply_blocks(document, REPORT_CONTENT)
    document.save(output)


def build_requirements_doc() -> None:
    output = REPORT_DIR / "需求说明书.docx"
    shutil.copyfile(BASE_DIR / "课程设计报告书.docx", output)
    document = Document(output)
    rename_headings(document, REQ_HEADING_RENAMES)
    set_cover(document, "需求说明书")
    apply_blocks(document, REQ_CONTENT)
    document.save(output)


def build_design_doc() -> None:
    output = REPORT_DIR / "设计文档.docx"
    shutil.copyfile(BASE_DIR / "设计文档.docx", output)
    document = Document(output)
    set_cover(document, "设计文档")
    apply_blocks(document, DESIGN_CONTENT)
    document.save(output)


def build_test_doc() -> None:
    output = REPORT_DIR / "测试文档.docx"
    shutil.copyfile(BASE_DIR / "测试文档.docx", output)
    document = Document(output)
    set_cover(document, "测试文档")
    apply_blocks(document, TEST_CONTENT)
    document.save(output)


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    build_report_doc()
    build_requirements_doc()
    build_design_doc()
    build_test_doc()


if __name__ == "__main__":
    main()
